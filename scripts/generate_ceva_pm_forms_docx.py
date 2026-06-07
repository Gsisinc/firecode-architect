"""
CEVA Logistics — ACS field PM & closeout forms (DOCX).

Generates **12** editable Word templates: Phases **A–I** (Forms 01–08) plus
**09** Network IP worksheet receipt, **10** Aux/IDF/fiber change notice,
**11** HID credential hold / release, **12** Daily toolbox coordination log.

Run: python scripts/generate_ceva_pm_forms_docx.py

For a **single handwriting-friendly PDF** (all twelve forms bound in order):
  python scripts/generate_ceva_pm_forms_pdf.py

Output folder (Word): ./CEVA_Closeout_Field_Forms_QIS/

Single PDF workbook: **`CEVA_Closeout_Field_Workbook_QIS.pdf`** (repository root).

"""

from __future__ import annotations

from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError as e:
    raise SystemExit("Install: pip install python-docx\n" + str(e)) from e

ROOT = Path(__file__).resolve().parent.parent
OUT_DIR = ROOT / "CEVA_Closeout_Field_Forms_QIS"

COLOR_NAVY = RGBColor(27, 58, 95)
COLOR_SLATE = RGBColor(89, 89, 89)
COLOR_WHITE = RGBColor(255, 255, 255)
FILL_HEADER = "1B3A5F"


def set_cell_fill(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def style_hdr_row(tbl, row_ix: int = 0) -> None:
    for cell in tbl.rows[row_ix].cells:
        set_cell_fill(cell, FILL_HEADER)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = COLOR_WHITE
                run.font.size = Pt(10)


def add_title_block(doc: Document, title: str, phase_tag: str, form_num: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("QUALITY INSTALLATION SYSTEMS\n")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_NAVY
    q = doc.add_paragraph()
    q.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s = q.add_run(f"{form_num}\n{title}")
    s.bold = True
    s.font.size = Pt(16)
    s.font.color.rgb = COLOR_NAVY
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    u = t.add_run(f"CEVA Logistics — ACS / Pro-Watch → Avigilon ACM + Mercury\nPhase reference: {phase_tag}")
    u.italic = True
    u.font.size = Pt(10)
    u.font.color.rgb = COLOR_SLATE
    doc.add_paragraph()


def add_site_header_table(doc: Document) -> None:
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    rows = [
        ("Site ID (e.g. US-MJB-04, PR-CAR-01)", ""),
        ("Facility / address", ""),
        ("CEVA Security Technology PM / Regional contact", ""),
        ("QIS lead tech / superintendent", ""),
        ("Form date (YYYY-MM-DD)", ""),
    ]
    for i, (lab, _) in enumerate(rows):
        tbl.rows[i].cells[0].text = lab
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        tbl.rows[i].cells[1].text = _
    doc.add_paragraph()


def sig_block_simple(doc: Document, roles: tuple[str, ...]) -> None:
    doc.add_paragraph().add_run("Signatures").bold = True
    tbl = doc.add_table(rows=len(roles) + 1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr[0].text = "Role"
    hdr[1].text = "Printed name"
    hdr[2].text = "Signature"
    hdr[3].text = "Date"
    style_hdr_row(tbl, 0)
    for i, role in enumerate(roles, start=1):
        tbl.rows[i].cells[0].text = role
        tbl.rows[i].cells[1].text = ""
        tbl.rows[i].cells[2].text = ""
        tbl.rows[i].cells[3].text = ""


def find_logo_for_forms() -> Path | None:
    for name in ("qis-logo.png", "qis-logo.jpg"):
        p = ROOT / "public" / "branding" / name
        if p.is_file():
            return p
    return None


def optional_logo(doc: Document) -> None:
    logo = find_logo_for_forms()
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(logo), width=Inches(2.2))
        doc.add_paragraph()


# --- Individual forms ------------------------------------------------------------


def form_01_mobilization(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(
        doc,
        "Site Mobilization & Pre-Job Checklist",
        "A / Pre-mobilization",
        "Form 01",
    )
    add_site_header_table(doc)
    doc.add_paragraph().add_run(
        "Use this checklist before cutover-critical work. Attach photos or email confirmations where noted."
    ).italic = True
    bullets = [
        "Site induction / EH&S acknowledgement complete (visitor log if required)",
        "CEVA badge / escort arrangements confirmed — names on gate list",
        "Site contact validated (facility + Security Technology stakeholder)",
        "MDF/IDF / headend electrical: 120 VAC termination within ~10 ft of enclosures (verify label + breaker ID)",
        "Delivery bay / forklift / pallet jack reserved for panel & cable offload",
        "Lift equipment reservation + operator certification on file",
        "OSHA-required PPE briefing for warehouse / lift work",
        "Credential approval: HID Signo / Corporate 1000 + 36-bit — **written CEVA release before issuing reader PO**",
        "Hardware receiving log started — compare packing slip vs BOM line tags (panels, readers, LSP trays)",
        "Network Support IP worksheet **requested / received** — attach copy when issued (cross-ref Form 09)",
        "Photo baseline: legacy ProWatch panels (door open, labeling, cabling dress)",
        "Confined coordination window agreed (quiet hours / dock blackout if applicable)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    h = tbl.rows[0].cells
    h[0].text = "#"
    h[1].text = "Shipment / pallet ID"
    h[2].text = "SKU / description"
    h[3].text = "Qty received"
    h[4].text = "Exception / shortage note"
    style_hdr_row(tbl, 0)
    for idx in range(14):
        r = tbl.add_row().cells
        r[0].text = str(idx + 1)
        r[1].text = r[2].text = r[3].text = r[4].text = ""

    sig_block_simple(
        doc,
        ("QIS site lead", "CEVA facility representative", "CEVA Security Technology (optional witness)"),
    )


def form_02_prowatch_decom(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(doc, "Pro-Watch Decommission Log", "B — Decommission", "Form 02")
    add_site_header_table(doc)
    doc.add_paragraph().add_run(
        "Remove and inventory legacy Pro-Watch equipment per CEVA rip-and-replace cable standard "
        "(log underground / abandoned conductors per survey)."
    ).italic = True

    tbl = doc.add_table(rows=1, cols=7)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    hdr_labels = ["#", "Legacy panel ID / location", "Cards/modules removed", "Cable cut tag color", "Spool recovered (ft)", "Inventory destination (dock / recycle)", "Photos Y/N"]
    for i, lab in enumerate(hdr_labels):
        hdr[i].text = lab
    style_hdr_row(tbl, 0)
    for _ in range(22):
        r = tbl.add_row().cells
        for i in range(7):
            r[i].text = ""

    doc.add_paragraph()
    doc.add_paragraph().add_run("Pull / abandonment notes").bold = True
    doc.add_paragraph("[Record trench IDs, conduit IDs, or abandoned homerun tags]")
    tbl2 = doc.add_table(rows=1, cols=4)
    tbl2.style = "Table Grid"
    h2 = tbl2.rows[0].cells
    h2[0].text = "Cable run ID"
    h2[1].text = "From → To"
    h2[2].text = "Disposition (removed / abandoned / splice retained)"
    h2[3].text = "Verified by"
    style_hdr_row(tbl2, 0)
    for _ in range(10):
        r = tbl2.add_row().cells
        for i in range(4):
            r[i].text = ""

    sig_block_simple(doc, ("QIS technician", "QIS superintendent",))


def form_03_rough_in_cable(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(doc, "Rough-In & Cable Pull Record", "C — Rough-in", "Form 03")
    add_site_header_table(doc)
    intro = doc.add_paragraph()
    r0 = intro.add_run(
        "Yellow plenum composite (standard ACS); Yellow plenum 22/6 SHLD for CI/CO. Each homerun ≤ 500 ft to "
        "headend/aux unless approved IDF (CEVA rule). Include "
    )
    r1 = intro.add_run("10 ft ")
    r2 = intro.add_run(
        "service loop allowance at the reader end per CEVA standard. Log billed footage from pull tickets; cite IDF/aux "
        "change notice (Form 10) if rule cannot be met."
    )
    for rr in (r0, r1, r2):
        rr.italic = True

    tbl = doc.add_table(rows=1, cols=11)
    tbl.style = "Table Grid"
    cols = [
        "#",
        "Door / portal ID",
        "Reader device ID",
        "Cable type (comp / SHLD)",
        "From rack / panel / aux",
        "To portal",
        "Billed ft (pull ticket)",
        "≤500 ft P/F",
        "Sweep / certify",
        "Label QC",
        "Notes / IDF ref",
    ]
    hdr = tbl.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c
    style_hdr_row(tbl, 0)
    for _ in range(28):
        rw = tbl.add_row().cells
        for i in range(11):
            rw[i].text = ""

    doc.add_paragraph()
    doc.add_paragraph("[Attach sweep / OLTS certification PDF if required by QA]")

    sig_block_simple(doc, ("Pull lead", "QC / foreman"))


def form_04_headend_comm(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(doc, "Headend Commissioning Record", "D / F — Install & power checks", "Form 04")
    add_site_header_table(doc)
    doc.add_paragraph().add_run(
        "LifeSafety Power dual-voltage (12 V controller rail + 24 V locking feeders). Mercury LP1502 / MR52 landing."
    ).italic = True

    tbl = doc.add_table(rows=1, cols=12)
    tbl.style = "Table Grid"
    cols = [
        "Panel/cluster",
        "LSP tag",
        "LP1502 S/N",
        "MR52 qty",
        "12 V measured",
        "24 V measured",
        "Breaker ID",
        "Ground bond OK",
        "ACM link / LED",
        "Static IP logged",
        "Subnet / GW / DNS per worksheet",
        "Punch notes",
    ]
    hdr = tbl.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c
    style_hdr_row(tbl, 0)
    for _ in range(14):
        r = tbl.add_row().cells
        for i in range(12):
            r[i].text = ""

    sig_block_simple(doc, ("Panel tech", "QIS PM / lead"))


def form_05_device_install(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(doc, "Device Installation & Termination Sign-Off", "E — Device install", "Form 05")
    add_site_header_table(doc)
    doc.add_paragraph().add_run(
        "Door-by-door: readers (HID Signo), REX, contacts, strikes / electrified hardware, sounders/strobes."
    ).italic = True

    tbl = doc.add_table(rows=1, cols=13)
    tbl.style = "Table Grid"
    cols = [
        "Door #",
        "CEVA door name (placeholder until naming workbook)",
        "Reader(s) terminated",
        "REX wired / aimed",
        "DPS / contact tested",
        "Strike / actuator",
        "Sounder/strobe",
        "ESL / cabling dress",
        "Torque / hardware QC",
        "Photos",
        "In ACM map",
        "SAT prep OK",
        "Notes",
    ]
    hdr = tbl.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c
    style_hdr_row(tbl, 0)
    for _ in range(36):
        r = tbl.add_row().cells
        for i in range(13):
            r[i].text = ""

    sig_block_simple(doc, ("Installer", "QC", "Facility witness (optional)"))


def form_06_acm_program(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(doc, "ACM Programming Record", "G — ACM programming", "Form 06")
    add_site_header_table(doc)
    doc.add_paragraph().add_run(
        "Controllers: static IP fields from Network Support. Browser UI only — no thick client installs. "
        "Document ACM license SKUs × qty for Server 1 + Server 2 (failover) including CEVA double-license doctrine — "
        "attach distributor PDF."
    ).italic = True

    tbl = doc.add_table(rows=1, cols=8)
    tbl.style = "Table Grid"
    cols = ["Controller MAC / ID", "IP", "Subnet", "Gateway", "DNS1", "DNS2", "Door count loaded", "License bundle ref"]
    for i, c in enumerate(cols):
        tbl.rows[0].cells[i].text = c
    style_hdr_row(tbl, 0)
    for _ in range(10):
        r = tbl.add_row().cells
        for i in range(8):
            r[i].text = ""

    doc.add_paragraph()
    doc.add_paragraph().add_run("Door naming — match CEVA official convention").bold = True
    t2 = doc.add_table(rows=1, cols=6)
    t2.style = "Table Grid"
    h2 = ["Door #", "Programmed name", "Reader IN", "Reader OUT", "Anti-passback segment", "Schedule ref"]
    for i, lab in enumerate(h2):
        t2.rows[0].cells[i].text = lab
    style_hdr_row(t2, 0)
    for _ in range(36):
        r = t2.add_row().cells
        for i in range(6):
            r[i].text = ""

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("VPN / remote support: ").bold = True
    p.add_run("Account issued to CEVA Security Team — expiry per policy (6 mo std / 1 yr max documented below).")

    tbl3 = doc.add_table(rows=3, cols=2)
    tbl3.style = "Table Grid"
    tbl3.rows[0].cells[0].text = "VPN credential ID"
    tbl3.rows[0].cells[1].text = ""
    tbl3.rows[1].cells[0].text = "Issue date"
    tbl3.rows[1].cells[1].text = ""
    tbl3.rows[2].cells[0].text = "Expiration date"
    tbl3.rows[2].cells[1].text = ""

    sig_block_simple(doc, ("QIS programmer", "CEVA validator (witness)"))


def form_07_sat(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(doc, "System Acceptance Test (SAT)", "H — Commissioning / SAT", "Form 07")
    add_site_header_table(doc)
    doc.add_paragraph().add_run(
        "Document each secured portal: free egress, electrified latch monitor, rex, ACS event to headend/log."
    ).italic = True

    tbl = doc.add_table(rows=1, cols=9)
    tbl.style = "Table Grid"
    cols = [
        "Door",
        "Card IN test",
        "Card OUT test",
        "Forced door",
        "REX bypass",
        "Strike monitor",
        "Alarm event good",
        "DVR/trace evidence",
        "Fail notes / punch",
    ]
    hdr = tbl.rows[0].cells
    for i, c in enumerate(cols):
        hdr[i].text = c
    style_hdr_row(tbl, 0)
    for _ in range(36):
        r = tbl.add_row().cells
        for i in range(9):
            r[i].text = ""

    doc.add_paragraph()
    doc.add_paragraph().add_run("Final acceptance signature — CEVA").bold = True
    doc.add_paragraph(
        "Per CEVA standards, final acceptance follows review by Security Technology PM **or** respective Regional Security Manager."
    )
    tbl2 = doc.add_table(rows=4, cols=2)
    tbl2.style = "Table Grid"
    rows_txt = [
        ("Print name (CEVA approving authority)", ""),
        ("Title", ""),
        ("Signature", ""),
        ("Acceptance date", ""),
    ]
    for i, (a, b) in enumerate(rows_txt):
        tbl2.rows[i].cells[0].text = a
        tbl2.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        tbl2.rows[i].cells[1].text = b

    sig_block_simple(doc, ("QIS commissioning lead",))


def form_08_closeout(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(doc, "Closeout & Turnover Package Checklist", "I — Closeout documentation", "Form 08")
    add_site_header_table(doc)
    doc.add_paragraph().add_run(
        "All as-built artifacts become property of CEVA Logistics upon acceptance. Maintain electronic + hard copy per contract."
    ).italic = True

    items = [
        "Marked-up floor plans showing panel + reader + cable IDs",
        "Rack elevations / LIU labeling (fiber segment if deployed)",
        "Panel schedules (voltages / breakers)",
        "IP matrix as-built matching Network worksheet",
        "ACM export / screenshot pack (hardware tree, door config summary)",
        "Test & inspection summary (Megger/FOL if applicable)",
        "Warranty registrar (panel, reader, PSU, LIC serials)",
        "Spare kit list handed to FM",
        "Training sign-in sheet(s) — if contracted",
        "Punch closure matrix with photos",
        "O&M excerpts + tech support escalation card",
        "Certificate of substantial completion readiness memo",
        "Cyber SSPS structured cable compliance statement (Cyber 270μm class)",
        "Cisco serial documentation — see supplement table immediately below",
        "FTP / portal upload confirmations (screenshots)",
        "VPN revocation / rollover requested post-SAT window",
        "Final invoice backup package (internal — optional attach)",
    ]
    for item in items:
        doc.add_paragraph(f"☐ {item}")

    doc.add_paragraph()
    doc.add_paragraph().add_run("Supplement — Cisco / optics serial capture (portal prep)").bold = True

    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for i, lab in enumerate(["#", "Device role", "Model", "Hostname if any", "Serial #", "Meraki/dashboard note ID"]):
        hdr[i].text = lab
    style_hdr_row(tbl, 0)
    for _ in range(24):
        r = tbl.add_row().cells
        for i in range(6):
            r[i].text = ""

    sig_block_simple(doc, ("QIS PM", "CEVA recipient (witness)"))


def form_09_network_receipt(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(doc, "Network Support IP Information Sheet — Receipt & Acknowledgement", "Coordination (pre-D/G)", "Form 09")
    add_site_header_table(doc)
    doc.add_paragraph().add_run(
        "Evidence that Vendor obtained IP worksheet **before** static addressing on controllers."
    ).italic = True
    bullets = [
        "Worksheet version / revision date",
        "Network Support issuing analyst",
        "VLAN IDs, subnet masks, gateway, DNS duo",
        "Reserved IPs per LP1502 / ACM server NIC (if directed)",
        "Change window for cutover ping test",
        "Meraki/dashboard tags if applicable",
        "Deviation log (IPs updated mid-project)",
    ]
    for b in bullets:
        doc.add_paragraph(b, style="List Bullet")

    tbl = doc.add_table(rows=8, cols=2)
    tbl.style = "Table Grid"
    fields = [
        "Worksheet received date",
        "Received by (QIS)",
        "Transmission method",
        "Conflicts flagged (Y/N + note)",
        "Resolution owner",
        "Closed date",
        "Attachment pointer (PDF name / mailbox ref)",
        "CEVA STPM acknowledgement (initials/date)",
    ]
    for i, lab in enumerate(fields):
        tbl.rows[i].cells[0].text = lab
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        tbl.rows[i].cells[1].text = ""

    sig_block_simple(doc, ("QIS superintendent",))


def form_10_idf_change(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(
        doc,
        "Auxiliary Panel / Fiber / IDF Routing Change Notice",
        "Triggered when homerun >500 ft or facilities direction",
        "Form 10",
    )
    add_site_header_table(doc)
    doc.add_paragraph().add_run(
        "Use when exceeding CEVA 500 ft envelope or relocating headend/aux per Site Leadership / STPM written approval."
    ).italic = True

    tbl = doc.add_table(rows=10, cols=2)
    tbl.style = "Table Grid"
    rows = [
        ("Trigger summary (survey ref)", ""),
        ("Affected portals / cable IDs", ""),
        ("Proposed aux panel designation", ""),
        ("Fiber SKU / OM3 strand count estimate", ""),
        ("LIU / patch strategy", ""),
        ("Added LP1502 / MR52 forecast", ""),
        ("Budget / CO reference #", ""),
        ("Photo log ID", ""),
        ("CEVA approving signature date", ""),
        ("Go-live reschedule impact", ""),
    ]
    for i, (lab, _) in enumerate(rows):
        tbl.rows[i].cells[0].text = lab
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        tbl.rows[i].cells[1].text = _

    sig_block_simple(doc, ("QIS PM / estimator", "CEVA STPM", "Site Leadership (if required)"))


def form_11_credential_hold_release(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(
        doc,
        "HID / Corporate 1000 Credential Hold & Release Authorization",
        "Before reader hardware PO issuance",
        "Form 11",
    )
    add_site_header_table(doc)
    doc.add_paragraph().add_run(
        "Formal release before placing reader orders aligns with CEVA Security Technology directives."
    ).italic = True
    doc.add_paragraph(
        "☐ Credential workbook / matrix tied to Corporate 1000 + 36-bit\n"
        "☐ Approver identified (Security Technology PM or delegate)\n"
        "☐ Hold lifted with written transmission reference\n"
        "☐ Reader ship schedule coordinated with receiving / docks"
    )
    tbl = doc.add_table(rows=5, cols=2)
    tbl.style = "Table Grid"
    labels = (
        ("Hold initiated (date / rationale)", ""),
        ("Written approver (print + title)", ""),
        ("Release timestamp + email / ticket ID", ""),
        ("Distributor quote linkage", ""),
        ("Phased delivery / substitutions", ""),
    )
    for i, (lv, _) in enumerate(labels):
        tbl.rows[i].cells[0].text = lv
        tbl.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        tbl.rows[i].cells[1].text = ""
    sig_block_simple(doc, ("QIS PM / purchasing",))


def form_12_daily_coordination(doc: Document) -> None:
    optional_logo(doc)
    add_title_block(
        doc,
        "Daily Field Coordination Log (Toolbox Digest)",
        "Optional — mobilization through SAT window",
        "Form 12",
    )
    add_site_header_table(doc)
    tbl = doc.add_table(rows=1, cols=6)
    tbl.style = "Table Grid"
    hdr_txt = (
        "Date",
        "Weather / lifts OK?",
        "Toolbox topic summary",
        "CEVA counterpart(s)",
        "Planned outages",
        "Risks / carryovers",
    )
    for i, txt in enumerate(hdr_txt):
        tbl.rows[0].cells[i].text = txt
    style_hdr_row(tbl, 0)
    for _ in range(20):
        rr = tbl.add_row().cells
        for i in range(6):
            rr[i].text = ""
    sig_block_simple(doc, ("QIS superintendent / foreman",))


FORM_BUILDERS = [
    ("01_Site_Mobilization_PreJob_Checklist.docx", form_01_mobilization),
    ("02_ProWatch_Decommission_Log.docx", form_02_prowatch_decom),
    ("03_RoughIn_Cable_Pull_Record.docx", form_03_rough_in_cable),
    ("04_Headend_Commissioning_Record.docx", form_04_headend_comm),
    ("05_Device_Installation_Termination_SignOff.docx", form_05_device_install),
    ("06_ACM_Programming_Record.docx", form_06_acm_program),
    ("07_SAT_System_Acceptance_Test.docx", form_07_sat),
    ("08_Closeout_Turnover_Checklist.docx", form_08_closeout),
    ("09_Network_IP_Worksheet_Receipt.docx", form_09_network_receipt),
    ("10_AUX_IDF_Fiber_Change_Notice.docx", form_10_idf_change),
    ("11_HID_Credential_Hold_Release.docx", form_11_credential_hold_release),
    ("12_Daily_Field_Coordination_Log.docx", form_12_daily_coordination),
]


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    for fname, builder in FORM_BUILDERS:
        d = Document()
        builder(d)
        out = OUT_DIR / fname
        d.save(str(out))
        print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
