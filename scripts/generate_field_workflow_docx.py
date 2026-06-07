# ruff: noqa: E501 — long procedure strings intentional
"""
Generate QIS Field Technician Workflow — CEVA Pro-Watch + Avigilon ACM.
Aligned to RFP-target schedule; checklist per phase.

Run: python scripts/generate_field_workflow_docx.py
Logo: public/branding/qis-logo.png | .jpg
Output: CEVA_ACS_Field_Workflow_QIS.docx
"""

from __future__ import annotations

import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
_SCRIPT_DIR = Path(__file__).resolve().parent
if str(_SCRIPT_DIR) not in sys.path:
    sys.path.insert(0, str(_SCRIPT_DIR))

from workflow_phase_specs import PHASE_SPECS, PhaseSpec  # noqa: E402

OUT_PATH = ROOT / "CEVA_ACS_Field_Workflow_QIS.docx"

COLOR_NAVY = RGBColor(27, 58, 95)
COLOR_ACCENT = RGBColor(0, 112, 192)
COLOR_SLATE = RGBColor(89, 89, 89)
COLOR_WHITE = RGBColor(255, 255, 255)
FILL_HEADER = "1B3A5F"
FILL_SUBTLE = "E8EEF4"
FILL_HOLD = "FFF4CC"
FILL_TIMELINE = "D6E8F5"
FILL_OK = "E8F4E8"

COMPANY = {
    "legal_name": "Quality Installation Systems",
    "street": "2 Robert Browning",
    "city_st_zip": "Knoxville, TN 37932",
    "phone": "865-774-4444",
    "email": "Robert@qualityinstallationsystems.com",
}

DOC_VERSION = "2.0"
RFP_TARGET_COMPLETION = "June 25, 2026"
AWARD_NOTIFICATION_RFP = "May 1, 2026 (confirm actual Notice to Proceed / NTP date per PO)"


def find_logo_path() -> Path | None:
    for name in ("qis-logo.png", "qis-logo.jpg"):
        p = ROOT / "public" / "branding" / name
        if p.is_file():
            return p
    return None


def set_fill(cell, hx: str) -> None:
    tc = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), hx)
    sh.set(qn("w:val"), "clear")
    tc.append(sh)


def hdr_row(tbl, ri: int = 0) -> None:
    for cell in tbl.rows[ri].cells:
        set_fill(cell, FILL_HEADER)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.font.color.rgb = COLOR_WHITE
                r.bold = True


def subtle_row(tbl, ri: int) -> None:
    for cell in tbl.rows[ri].cells:
        set_fill(cell, FILL_SUBTLE)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


def style_ok_banner(cell) -> None:
    set_fill(cell, FILL_OK)
    for p in list(cell.paragraphs):
        p._element.getparent().remove(p._element)
    rp = cell.add_paragraph()
    rp.paragraph_format.space_before = Pt(4)
    rp.paragraph_format.space_after = Pt(4)


def add_h(doc: Document, t: str, level: int = 1) -> None:
    doc.add_heading(t, level=level)


def phase_banner(doc: Document, num: str, title: str) -> None:
    b = doc.add_table(rows=1, cols=1)
    c = b.rows[0].cells[0]
    set_fill(c, FILL_HEADER)
    for p in list(c.paragraphs):
        p._element.getparent().remove(p._element)
    pr = c.add_paragraph()
    pr.paragraph_format.space_before = Pt(7)
    pr.paragraph_format.space_after = Pt(7)
    r = pr.add_run(f"PHASE {num}  —  {title.upper()}")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_WHITE


def timeline_callout(doc: Document, weeks: str) -> None:
    t = doc.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Target window\n(from site NTP)"
    t.rows[0].cells[1].text = weeks
    set_fill(t.rows[0].cells[0], FILL_TIMELINE)
    set_fill(t.rows[0].cells[1], FILL_TIMELINE)
    for cell in t.rows[0].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLOR_NAVY


def phase_checklist(doc: Document, pid: str) -> None:
    add_h(doc, f"Phase {pid} completion checklist — sign off each row before leaving phase", level=3)
    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    h = t.rows[0].cells
    cols = ["Done", "#", "Action / criterion", "How you verify", "Proof / record"]
    for i, col in enumerate(cols):
        h[i].text = col
    hdr_row(t, 0)


def add_check_rows(t, rows: tuple[tuple[str, str, str], ...]) -> None:
    for n, (act, verify, proof) in enumerate(rows, start=1):
        r = t.add_row().cells
        r[0].text = "☐"
        r[1].text = str(n)
        r[2].text = act
        r[3].text = verify
        r[4].text = proof


def gate_row(doc: Document, txt: str) -> None:
    tb = doc.add_table(rows=1, cols=1)
    c = tb.rows[0].cells[0]
    set_fill(c, FILL_HOLD)
    c.text = f"HOLD — {txt}"
    for p in c.paragraphs:
        for r in p.runs:
            r.bold = True


def render_phase(doc: Document, ps: PhaseSpec) -> None:
    doc.add_page_break()
    phase_banner(doc, ps.num, ps.title)
    timeline_callout(doc, ps.week_window)

    ip = doc.add_paragraph(ps.intro)
    for r in ip.runs:
        r.italic = True
        r.font.color.rgb = COLOR_SLATE

    add_h(doc, "Objectives — must be true before exiting this phase", level=3)
    for o in ps.objectives:
        doc.add_paragraph(o, style="List Bullet")

    add_h(doc, "Detailed procedure (execute in order unless Lead reprioritizes for parallel sites)", level=3)
    for i, proc in enumerate(ps.procedures, start=1):
        doc.add_paragraph(f"{ps.num}.{i:02d}  {proc}", style="List Number")

    phase_checklist(doc, ps.num)
    # last table added is checklist — get it from doc.tables[-1]? fragile
    chk = doc.tables[-1]
    add_check_rows(chk, ps.checklist)
    exit_row = chk.add_row().cells
    exit_row[0].text = "☐"
    exit_row[1].text = "✓"
    exit_row[2].text = "Lead sign-off: phase complete (initials / date)"
    exit_row[3].text = "Walk with Lead; no open hazards; iCare note"
    exit_row[4].text = "WO comment + photo index"
    subtle_row(chk, len(chk.rows) - 1)

    if ps.hold_gate_after:
        doc.add_paragraph()
        gate_row(doc, ps.hold_gate_after)


def cover(doc: Document) -> None:
    doc.core_properties.title = "Field Workflow — CEVA ACS"
    doc.core_properties.author = COMPANY["legal_name"]
    lg = find_logo_path()
    if lg:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(lg), width=Inches(2.58))
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        x = p.add_run("Logo: public/branding/qis-logo.png (or .jpg) — re-run to embed.")
        x.italic = True
        x.font.color.rgb = COLOR_SLATE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FIELD OPERATIONS WORKFLOW")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = COLOR_ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Enterprise access control rollout")
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = COLOR_NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"Quality Installation Systems  •  {COMPANY['phone']}  •  {COMPANY['email']}\n").font.color.rgb = (
        COLOR_SLATE
    )
    doc.add_paragraph()

    stripe = doc.add_table(rows=1, cols=1)
    sc = stripe.rows[0].cells[0]
    set_fill(sc, FILL_TIMELINE)
    for p in list(sc.paragraphs):
        p._element.getparent().remove(p._element)
    pr = sc.add_paragraph()
    pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pr.add_run(
        f"Program milestone (RFP schedule reference): complete by {RFP_TARGET_COMPLETION}\n"
        f"Award / NTP reference: {AWARD_NOTIFICATION_RFP}\n"
        "Weekly bands below count from your site Notice to Proceed — adjust with PM if dates slip."
    )
    rr.font.bold = True
    rr.font.color.rgb = COLOR_NAVY
    rr.font.size = Pt(11)

    ctrl = doc.add_table(rows=4, cols=2)
    ctrl.style = "Table Grid"
    ctrl.rows[0].cells[0].text = "Control"
    ctrl.rows[0].cells[1].text = "Value"
    hdr_row(ctrl, 0)
    dv = DOC_VERSION  # shorten line
    for i, pair in enumerate(
        [
            ("Document revision", dv),
            ("Printed", date.today().strftime("%Y-%m-%d")),
            ("Applies to", "CEVA Pro-Watch decommission + Avigilon ACM / Mercury (per WO site ID)"),
        ],
        start=1,
    ):
        ctrl.rows[i].cells[0].text = pair[0]
        ctrl.rows[i].cells[1].text = pair[1]


def master_timeline(doc: Document) -> None:
    doc.add_page_break()
    add_h(doc, "Master timeline — phases vs calendar (from Notice to Proceed)", level=1)
    doc.add_paragraph(
        f"Contracts often bind a substantial completion (~{RFP_TARGET_COMPLETION} on the published RFP). "
        "This table is a pacing guide for ONE site executing with a typical 2-tech crew — compress or overlap only when PM schedules it. "
        "Two Mount Juliet buildings in parallel requires staggered lifts and duplicated material drops."
    )
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    heads = ["Phase", "Target weeks from NTP", "Primary outputs", "Exit gate"]
    for i, h in enumerate(heads):
        t.rows[0].cells[i].text = h
    hdr_row(t, 0)
    grid = [
        ("3 Day-zero / safety / walk", "1", "Photos, MDF/ACP walk, iCare alive", "Security + Lead agree door plan"),
        ("4 Demo Pro-Watch", "1–3", "Field clear, rack stripped, inventories", "No unsecured doors overnight"),
        ("5 Pathway / rough-in", "2–5", "J-hooks / stubs / service loops queued", ">500 ft runs flagged"),
        ("6 Cable pulls + test", "3–6", "Labeled homeruns, continuity done", "No ceiling closed without QA"),
        ("7 Headend install", "4–7", "LSP / LP1502 / MR52 landed", "DC stable, labeling 100 %"),
        ("8 Doors & devices", "4–8 (overlap)", "Readers, strikes, IDS/DPDT, IN/OUT", "Each opening punched before SAT"),
        ("9 Power-on / ladder checks", "5–8", "Reader/door proofs", "Punch tracker current"),
        ("10 ACM programming / IP", "6–9", "DB matches as-built numbering", "IT sheet matched"),
        ("11 SAT + punch closure", "7–10", "Witness tests done", "iCare SAT language"),
        ("12 Documentation turnover", "8–11+", "Drawings / panel schedules", "PM package complete"),
    ]
    for row in grid:
        c = t.add_row().cells
        for i, v in enumerate(row):
            c[i].text = v

    gb = doc.add_table(rows=1, cols=1)
    gx = gb.rows[0].cells[0]
    style_ok_banner(gx)
    ru = gx.add_paragraph().add_run(
        "If time from Notice to Proceed to turnover is shorter than these bands — add crew capacity, lengthen days, shrink scope phases, "
        "or negotiate phased SAT in writing."
    )
    ru.bold = True
    ru.font.color.rgb = COLOR_NAVY
    ru.font.size = Pt(10)


def how_to_use(doc: Document) -> None:
    doc.add_page_break()
    add_h(doc, "How to execute this playbook", level=1)
    intro = doc.add_paragraph()
    rr = intro.add_run(
        "Each execution phase opens with a WEEK WINDOW counting from Notice to Proceed. Complete the procedural steps in order unless "
        "the Lead adjusts for parallel installs (two Tennessee sites, etc.). "
        "Always finish by filling the PHASE CHECKLIST and obtaining Lead initials."
    )
    rr.font.color.rgb = COLOR_SLATE

    tips = doc.add_table(rows=1, cols=2)
    tips.style = "Table Grid"
    tips.rows[0].cells[0].text = "Symbol / rule"
    tips.rows[0].cells[1].text = "Meaning"
    hdr_row(tips, 0)
    for lab, txt in [
        ("☐ checkboxes", "Print-friendly; fill in Acrobat or replicate in OneNote."),
        ("HOLD bar (amber)", "Hard stop until condition cleared — escalate to Lead/PM same day."),
        ("Phase banner (navy)", "Scope boundary — do not shortcut into next phase."),
        ("iCare / WO hygiene", "Status + photos anytime customer policy requires."),
    ]:
        row = tips.add_row().cells
        row[0].text = lab
        row[1].text = txt


def section_roles(doc: Document) -> None:
    doc.add_page_break()
    add_h(doc, "Roles & stop-work (read once per job)", level=1)
    sw = doc.add_table(rows=1, cols=2)
    sw.style = "Table Grid"
    sw.rows[0].cells[0].text = "Signal"
    sw.rows[0].cells[1].text = "Mandatory response"
    hdr_row(sw, 0)
    for cond, resp in [
        ("No NIC IP sheet aligned to racks", "No panel on LAN."),
        ("Any homerun modeled >500 ft without IDF", "Stop pull."),
        ("Fail-safe electrified gear needs FA relay", "Schedule FA subcontract; no improvised jumpers."),
        ("Scope creep vs WO (gates/intercom additions)", "PM change order."),
    ]:
        r = sw.add_row().cells
        r[0].text = cond
        r[1].text = resp


def section_prep(doc: Document) -> None:
    doc.add_page_break()
    add_h(doc, "Pre-trip — PM packs / truck stock", level=2)
    t = doc.add_table(rows=1, cols=3)
    t.style = "Table Grid"
    h = ["☐", "Item", "Remark"]
    for i, head in enumerate(h):
        t.rows[0].cells[i].text = head
    hdr_row(t, 0)
    for row in [
        ("WO + POC contacts + badges", ""),
        ("Latest CAD/PDF drawings + markup pen", ""),
        ("IP sheet DNS/gateway finalized", ""),
        ("Credential format lock (HID corp1000)", ""),
        ("Lift POC + delivery ETA", ""),
        ("Label rolls + toner + ladder cert", ""),
    ]:
        c = t.add_row().cells
        c[0].text = "☐"
        c[1].text = row[0]
        c[2].text = row[1]


def rollback(doc: Document) -> None:
    doc.add_page_break()
    add_h(doc, "Emergency / rollback shorthand", level=1)
    z = doc.add_table(rows=1, cols=2)
    z.style = "Table Grid"
    z.rows[0].cells[0].text = "Risk"
    z.rows[0].cells[1].text = "Stabilize"
    hdr_row(z, 0)
    for a, b in [
        ("Cannot secure door electrically", "Secure mechanically + guard + escalate."),
        ("Smoke or heat near cable pull", "Stop. Evacuate per alarm. Supervisor call."),
        ("Panel brownout spikes", "Remove load steps; diagnose ground faults before re-attach."),
    ]:
        r = z.add_row().cells
        r[0].text = a
        r[1].text = b


def build() -> None:
    doc = Document()
    cover(doc)
    master_timeline(doc)
    how_to_use(doc)
    section_roles(doc)
    section_prep(doc)
    for ps in PHASE_SPECS:
        render_phase(doc, ps)
    rollback(doc)

    doc.add_page_break()
    add_h(doc, "End / revision footer", level=1)
    doc.add_paragraph(
        f"Quality Installation Systems — Field playbook v{DOC_VERSION} — printed {date.today().isoformat()}"
    )
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote: {OUT_PATH}")


if __name__ == "__main__":
    build()