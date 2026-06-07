"""
CEVA ACS — Addon package estimates (NOT in base bid).
Separate section per site: gates, turnstiles/ADA, intercoms, network/Cisco/fiber contingency.

Run:
  python scripts/generate_ceva_addon_estimates_docx.py                         # TN two-site + rollup
  python scripts/generate_ceva_addon_estimates_docx.py --site US-LUI-02        # Louisville KY only

Outputs:
  CEVA_Addon_Packages_Estimate_QIS.docx               (TN program default)
  CEVA_Addon_Packages_Estimate_QIS_US-LUI-02.docx     (Kentucky)

Logo: public/branding/qis-logo.png | .jpg
"""

from __future__ import annotations

import argparse
from dataclasses import dataclass
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = ROOT / "CEVA_Addon_Packages_Estimate_QIS.docx"
OUT_PATH_LUI02 = ROOT / "CEVA_Addon_Packages_Estimate_QIS_US-LUI-02.docx"

LABOR_RATE = 155.00
PM_PCT = 0.12
MARGIN_PCT = 0.25

COLOR_NAVY = RGBColor(27, 58, 95)
COLOR_ACCENT = RGBColor(0, 112, 192)
COLOR_SLATE = RGBColor(89, 89, 89)
COLOR_WHITE = RGBColor(255, 255, 255)
FILL_HEADER = "1B3A5F"
FILL_SUBTLE = "E8EEF4"

COMPANY = {
    "legal_name": "Quality Installation Systems",
    "street": "2 Robert Browning",
    "city_st_zip": "Knoxville, TN 37932",
    "phone": "865-774-4444",
    "email": "Robert@qualityinstallationsystems.com",
}


@dataclass
class AddonLine:
    code: str
    title: str
    assumption: str
    hours_low: float
    hours_mid: float
    hours_high: float
    material_note: str


def money(n: float) -> str:
    return f"${n:,.2f}"


def labor_sell(hours: float) -> tuple[float, float]:
    """Returns (sell_total, direct_labor_only)."""
    direct = hours * LABOR_RATE
    pm = direct * PM_PCT
    pre = direct + pm
    mrg = pre * MARGIN_PCT
    return pre + mrg, direct


def find_logo() -> Path | None:
    for n in ("qis-logo.png", "qis-logo.jpg"):
        p = ROOT / "public" / "branding" / n
        if p.is_file():
            return p
    return None


def set_fill(cell, hx: str) -> None:
    tc = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:fill"), hx)
    sh.set(qn("w:val"), "clear")
    tc.append(sh)


def hdr_style(tbl, ri: int = 0) -> None:
    for cell in tbl.rows[ri].cells:
        set_fill(cell, FILL_HEADER)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = COLOR_WHITE


def subtle_row(tbl, ri: int) -> None:
    for cell in tbl.rows[ri].cells:
        set_fill(cell, FILL_SUBTLE)
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True


ADDONS_MJB04: list[AddonLine] = [
    AddonLine(
        "PKG-B",
        "Vehicle gates package (dual-height pedestals / readers)",
        "Planning basis: truck gate pedestal count from walk-through (4 pedestal faces); inbound+outbound vehicle lanes "
        "per CEVA program (Wright-style); excludes gate operators, loops, barriers, civil (per typical CEVA language). "
        "Underground-rated cable pulls + terminations + interface testing.",
        80,
        120,
        175,
        "Pedestals, readers, keypad-upper/reader-lower, cable, conduit (excluded vendor hardware per RFP clarification). Quote separate.",
    ),
    AddonLine(
        "PKG-C",
        "Full-height turnstile + ADA gate (Alvarado-class integration)",
        "Planning basis: 1 turnstile + 1 ADA path from site notes; inbound/outbound readers, crash-bar on ADA, DPS, "
        "sounders if IN/OUT policy applies — confirm on drawings.",
        28,
        48,
        72,
        "Turnstile / ADA equipment by others or GC package often — ACS integration labor + minor materials only unless BOM states otherwise.",
    ),
    AddonLine(
        "PKG-A",
        "Intercom subsystem (Axis / 2N substations + answering stations + ACC)",
        "Planning basis: dual-height gate package often drives 8 substation-mount points + 4 desk masters (program call-all-masters); "
        "Green Cat6 homeruns + composite spare pairs for triggers; ACC video-channel setup. Adjust counts when SOW finalized.",
        72,
        105,
        145,
        "Intercom hardware, mounts, UPS at masters — distributor BOM; excludes ACC server licensing/storage beyond scope.",
    ),
    AddonLine(
        "PKG-D",
        "Network augmentation — Cisco licensing, optics, fiber jumps, hardened edge (vendor-responsible per CEVA IT matrix)",
        "QIS holds the PO when CEVA directs Network Support handoffs: serial documentation, patch/fiber execution, "
        "hardened-switch entitlements, optics kitting. Hours flex with switch counts + IDF outcomes.",
        6,
        18,
        36,
        "LIC-C9300 / LIC-9300L / IE entitlement + GLC-SX-MMD / SFP-10G-SR + OM3 LC jumpers — all non-returnable; "
        "finalize counts with Network Support before release.",
    ),
]

ADDONS_MJB02: list[AddonLine] = [
    AddonLine(
        "PKG-B",
        "Vehicle gates package (dual-height pedestals / readers)",
        "Same gate methodology as sibling site; MJB-02 notes include 4 truck gate pedestals and heavier dock/ramp footprint — "
        "labor midpoint slightly higher than MJB-04.",
        90,
        135,
        200,
        "Same exclusions: gate operators/hardware excluded pending RFP carve-out confirmation.",
    ),
    AddonLine(
        "PKG-C",
        "Full-height turnstile + ADA gate integration",
        "Planning basis: 1 turnstile + ADA from walk-through notes.",
        28,
        48,
        72,
        "Equipment sourcing same as Site 01.",
    ),
    AddonLine(
        "PKG-A",
        "Intercom subsystem (Axis / 2N + ACC video channels)",
        "Slightly denser dock/gate clustering in notes — midpoint hours + programming higher than Site 01 until drawing-based count.",
        85,
        125,
        170,
        "Hardware BOM TBD — confirm substation count per finalized gate plan.",
    ),
    AddonLine(
        "PKG-D",
        "Network augmentation — Cisco optics + fiber + hardened edge (vendor-held per CEVA IT direction)",
        "Primary site still gets the larger Network Support touch plan; hours cover documentation, jumpering, portal entry "
        "prep, cutover windows with STPM.",
        10,
        22,
        45,
        "Non-returnable LIC + SFP stack; align quantities to same reference SKUs as mainland program.",
    ),
]

# Louisville KY — CL Warehouse (same PKG codes as program; bands sized for KY footprint pending survey).
ADDONS_LUI02: list[AddonLine] = [
    AddonLine(
        "PKG-B",
        "Vehicle gates package (dual-height pedestals / readers)",
        "Planning basis: 5200 Interchange CL warehouse — verify truck lane count and pedestal layout on drawings; "
        "typical dual-height pedestal program (inbound/outbound lanes) aligns with sibling CEVA sites until site walk adjusts counts. "
        "Excludes gate operators, loops, barriers, civil. Underground-rated pulls + terminations + interface testing.",
        78,
        115,
        168,
        "Pedestals, readers, keypad/reader stacks, conduit — quote distributor once lane plan is sealed. Operators excluded pending RFP.",
    ),
    AddonLine(
        "PKG-C",
        "Full-height turnstile + ADA gate (Alvarado-class integration)",
        "Planning basis: 1 turnstile + 1 ADA path assumed until egress sheets confirm; inbound/outbound readers, crash-bar on ADA "
        "(DPS / sounders if IN/OUT policy applies).",
        28,
        48,
        72,
        "Turnstile / ADA equipment frequently GC or OEM package — ACS integration labor + minor materials unless BOM states otherwise.",
    ),
    AddonLine(
        "PKG-A",
        "Intercom subsystem (Axis / 2N substations + answering stations + ACC)",
        "Planning basis: Louisville gate/dock clustering — midpoint slightly below larger Mount Juliet footprints; "
        "adjust substation counts when finalized along PKG-B perimeter. "
        "Green Cat6 homeruns + composite spares for triggers; ACC channel setup scope as awarded.",
        65,
        98,
        138,
        "Intercom hardware / UPS — distributor BOM; excludes ACC licensing/storage beyond negotiated SOW.",
    ),
    AddonLine(
        "PKG-D",
        "Network augmentation — Cisco optics + fiber kitting (still vendor-directed when auxiliary budget isolates PKG-D)",
        "Even when PKG-D splits out for bookkeeping, labor band keeps QIS on point for patching, jumper install, hardened "
        "edge prep per CEVA Network Support playbook.",
        6,
        16,
        34,
        "9300 LIC classes + optics non-returnable; sync counts with HQ switch workbook.",
    ),
]


def site_section(doc: Document, site_id: str, site_name: str, address: str, lines: list[AddonLine]) -> None:
    doc.add_page_break()
    doc.add_heading(f"{site_id} — addon packages (exclude from base bid)", level=1)
    p = doc.add_paragraph()
    r = p.add_run(f"{site_name}\n{address}")
    r.bold = True
    ip = doc.add_paragraph()
    ir = ip.add_run(
        "These appendix packages isolate perimeter subsystems (**PKG-A/B/C**) and incremental Cisco/optics labor beyond the "
        "vendor-owned placeholders already seeded in each site’s BOM. Hours are planning bands; midpoints drive pricing below. "
        f"Labor math matches core proposal (${LABOR_RATE:.2f}/hr direct + {PM_PCT:.0%} PM dollars + "
        f"{MARGIN_PCT:.0%} labor margin)."
    )
    ir.italic = True
    ir.font.color.rgb = COLOR_SLATE

    t = doc.add_table(rows=1, cols=6)
    t.style = "Table Grid"
    h = t.rows[0].cells
    h[0].text = "Code"
    h[1].text = "Addon"
    h[2].text = "Hrs (low–mid–high)"
    h[3].text = "Mid labor sell ⓘ"
    h[4].text = "Material (plan)"
    h[5].text = "Key assumption"
    hdr_style(t, 0)

    total_mid_hours = 0.0
    total_mid_sell = 0.0

    for ln in lines:
        sell_mid, _ = labor_sell(ln.hours_mid)
        total_mid_hours += ln.hours_mid
        total_mid_sell += sell_mid
        row = t.add_row().cells
        row[0].text = ln.code
        row[1].text = ln.title
        row[2].text = f"{ln.hours_low:.0f} / {ln.hours_mid:.0f} / {ln.hours_high:.0f}"
        row[3].text = money(sell_mid)
        row[4].text = ln.material_note
        row[5].text = ln.assumption

    sr = t.add_row().cells
    sr[0].text = "TOTAL (midpoint hrs)"
    sr[1].text = "Σ packages"
    sr[2].text = f"{total_mid_hours:.0f} hrs (mid components only)"
    sr[3].text = money(total_mid_sell)
    sr[4].text = "Mid labor sells summed row-by-row (PM/margin embedded per line)."
    sr[5].text = "Material: quote distributor per final SOW."
    subtle_row(t, len(t.rows) - 1)

    note = doc.add_paragraph()
    nr = note.add_run(
        f"ⓘ Mid labor sell = ({LABOR_RATE}/hr × mid hours) + {PM_PCT:.0%} PM + {MARGIN_PCT:.0%} margin chain. "
        "Low/high hours scale roughly proportionally — PM/margin recomputed linearly."
    )
    nr.italic = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = COLOR_SLATE


def cover(doc: Document, subtitle: str | None = None) -> None:
    doc.core_properties.title = "CEVA ADD-ON PACKAGE ESTIMATES"

    if subtitle is None:
        subtitle = (
            "Perimeter specialty + incremental Network overlays\n"
            "Base modernization proposals already embed vendor-held Cisco/optics placeholders."
        )

    lg = find_logo()
    if lg:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(lg), width=Inches(2.4))
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("OPTION / ADD-ON ESTIMATES")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = COLOR_ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = COLOR_NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f"{COMPANY['legal_name']} • {COMPANY['phone']} • {COMPANY['email']}\n"
        f"{COMPANY['street']}, {COMPANY['city_st_zip']}\n"
        f"Schedule of options • {date.today().strftime('%B %d, %Y')}\n"
    )


def combined_summary(doc: Document) -> None:
    doc.add_page_break()
    doc.add_heading("Program summary — both sites (midpoint)", level=1)
    t = doc.add_table(rows=4, cols=3)
    t.style = "Table Grid"
    t.rows[0].cells[0].text = "Site"
    t.rows[0].cells[1].text = "Mid planning hours (all addons)"
    t.rows[0].cells[2].text = "Mid labor sell (sum of lines)"
    hdr_style(t, 0)

    def sum_site(lines: list[AddonLine]) -> tuple[float, float]:
        th = sum(x.hours_mid for x in lines)
        ts = sum(labor_sell(x.hours_mid)[0] for x in lines)
        return th, ts

    h04, s04 = sum_site(ADDONS_MJB04)
    h02, s02 = sum_site(ADDONS_MJB02)
    t.rows[1].cells[0].text = "US-MJB-04"
    t.rows[1].cells[1].text = f"{h04:.0f}"
    t.rows[1].cells[2].text = money(s04)
    t.rows[2].cells[0].text = "US-MJB-02"
    t.rows[2].cells[1].text = f"{h02:.0f}"
    t.rows[2].cells[2].text = money(s02)
    t.rows[3].cells[0].text = "BOTH — reference only"
    t.rows[3].cells[1].text = f"{h04 + h02:.0f}"
    t.rows[3].cells[2].text = money(s04 + s02)
    subtle_row(t, 3)

    p_mat = doc.add_paragraph()
    p_mat.add_run(
        "The core proposal already seeds Cisco / optics / rack placeholders on each site BOM — these ADD-ON lines only "
        "capture **additional** SKU depth or SAT-intensive cutovers if Network Support expands quantities beyond baseline."
    )


def build() -> None:
    doc = Document()
    cover(doc)
    doc.add_paragraph()

    intr = doc.add_paragraph()
    ir = intr.add_run(
        "Purpose: PKG-A/B/C segregate perimeter / voice-heavy packages; PKG-D only scales labor when Network-directed "
        "Cisco/optics workloads exceed what the baseline BOM placeholders already quantify for QIS procurement."
    )
    ir.font.color.rgb = COLOR_SLATE

    site_section(doc, "US-MJB-04", "G&R Warehouse", "12014 Volunteer Blvd, Mount Juliet, TN", ADDONS_MJB04)
    site_section(doc, "US-MJB-02", "CL Warehouse (Building 3)", "12008 Volunteer Blvd, Building 3, Mount Juliet, TN", ADDONS_MJB02)
    combined_summary(doc)

    doc.add_paragraph()
    doc.add_heading("Disclaimer", level=2)
    doc.add_paragraph(
        "Indicative planning hours only until gate lane counts, intercom substation quantities, Turnstile model sheets, "
        "and Cisco port maps are finalized. Selecting any option triggers a distributor material quote and coordinated SAT plan."
    )

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote: {OUT_PATH}")


def build_louisville() -> None:
    doc = Document()
    cover(
        doc,
        subtitle="Option & overflow schedules (overlay US-LUI-02 base modernization)\n"
        "CEVA Logistics — Louisville KY",
    )
    doc.add_paragraph()

    intr = doc.add_paragraph()
    ir = intr.add_run(
        "Purpose: PKG-A/B/C segregate perimeter / voice-heavy packages; PKG-D only scales incremental Network-driven "
        "Cisco or fiber cutover labor beyond the BOM placeholders seeded in Louisville’s modernization proposal."
    )
    ir.font.color.rgb = COLOR_SLATE

    site_section(
        doc,
        "US-LUI-02",
        "CL Warehouse",
        "5200 Interchange Drive, Louisville, KY 40229",
        ADDONS_LUI02,
    )

    doc.add_paragraph()
    doc.add_heading("Disclaimer", level=2)
    doc.add_paragraph(
        "Indicative planning hours until gate lanes, intercom substation quantities, turnstile sheets, "
        "and Cisco port maps are finalized. Selecting any option triggers distributor material quoting and SAT coordination."
    )

    OUT_PATH_LUI02.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH_LUI02))
    print(f"Wrote: {OUT_PATH_LUI02}")


def main() -> None:
    p = argparse.ArgumentParser(description="CEVA ACS addon estimates (DOCX).")
    p.add_argument(
        "--site",
        choices=("TN_PAIR", "US-LUI-02"),
        default="TN_PAIR",
        help="TN_PAIR = Mount Juliet sites + program rollup (default); US-LUI-02 = Louisville only.",
    )
    args = p.parse_args()
    if args.site == "US-LUI-02":
        build_louisville()
    else:
        build()


if __name__ == "__main__":
    main()
