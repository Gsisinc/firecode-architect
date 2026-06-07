"""
Generate CEVA ProWatch Decommission + Avigilon ACM bid (DOCX).

Run examples:
  python scripts/generate_ceva_bid_docx.py                         # TN two-site → CEVA_Proposal_QIS_TN-MJB-Pair.docx
  python scripts/generate_ceva_bid_docx.py --site US-LUI-02        # Louisville KY
  python scripts/generate_ceva_bid_docx.py --site PR-CAR-01        # Puerto Rico addendum site

Logo: place PNG at public/branding/qis-logo.png (optional).
"""

from __future__ import annotations

import argparse
import math
from datetime import date, timedelta
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError as e:
    raise SystemExit("Install: pip install python-docx\n" + str(e)) from e

# Brand accents (RGB) — QIS / professional proposal
COLOR_NAVY = RGBColor(27, 58, 95)  # #1B3A5F
COLOR_ACCENT = RGBColor(0, 112, 192)
COLOR_SLATE = RGBColor(89, 89, 89)
COLOR_WHITE = RGBColor(255, 255, 255)
FILL_HEADER = "1B3A5F"  # hex without # for Word shading
FILL_SUBTLE = "E8EEF4"

ROOT = Path(__file__).resolve().parent.parent
OUT_PATH = ROOT / "CEVA_Proposal_QIS_TN-MJB-Pair.docx"
OUT_PATH_LUI02 = ROOT / "CEVA_Proposal_QIS_US-LUI-02.docx"
OUT_PATH_PR_CAR01 = ROOT / "CEVA_Proposal_QIS_PR-CAR-01.docx"


LABOR_RATE = 155.00
PM_PCT = 0.12  # PM as % of direct labor dollars (same as 12% of hours × rate)
MARGIN_PCT = 0.25  # 25% midpoint of 20–30% on labor subtotal and on material base
TAX_MATERIAL_PCT = 0.0925  # default / TN-style placeholder — override per-site with sales_tax_pct
BOND_PCT = 0.015  # performance/payment bond allowance — verify with surety
OCIP_PCT = 0.005  # contractor OCIP/wrap fee allowance when required — often $0 if owner-provided

# --- Cable takeoff (CEVA: 500 ft max field device → headend unless aux / IDF) ---
# Estimate: per homerun, (horizontal + vertical path + slack guess) = avg_homerun_ft, then add
# service loop at field end, then apply waste %. **Each assumed avg must stay ≤ max_run_ft** until
# survey; if any opening exceeds max, add **IDF / auxiliary panel / fiber** (change order).
SERVICE_LOOP_FT = 10.0
WASTE_PCT = 0.10
MAX_RUN_FT = 500.0
COMPOSITE_PLENUM_PER_FT = 1.22  # indicative — replace with distributor
SHIELDED_22_6_PER_FT = 2.95  # IN/OUT homeruns — indicative

COMPANY = {
    "legal_name": "Quality Installation Systems",
    "street": "2 Robert Browning",
    "city_st_zip": "Knoxville, TN 37932",
    "phone": "865-774-4444",
    "email": "Robert@qualityinstallationsystems.com",
    "website": "",
}

SIGNER_NAME = "Robert Browning"
SIGNER_TITLE = "Authorized Signer"
PROPOSAL_VALID_DAYS = 15

SITE1 = {
    "id": "US-MJB-04",
    "name": "G&R Warehouse",
    "address": "12014 Volunteer Blvd, Mount Juliet, TN",
    "readers_bid": 19,
    "lp1502": 7,
    "phases": [
        ("A — Pre-mobilization (setup, coordination, prerequisites)", 20),
        ("B — Pro-Watch decommission (remove, cable, inventory, light patch)", 76),
        ("C — Rough-in (pathway, homeruns, service loops, headend rough-in)", 96),
        ("D — Headend install (LSP / LP1502 / MR52 landing & checkout)", 70),
        ("E — Device install (readers, terminations, door interfaces — spread est.)", 70),
        ("F — Power-on checks & comms troubleshooting", 18),
        ("G — ACM programming (hardware map, doors, schedules per contract)", 50),
        ("H — Commissioning / SAT & punch", 24),
        ("I — Closeout documentation (as-builts, panel schedules, turnover)", 40),
    ],
    # Indicative market-cost BOM core (pre-margin); cable + lift appended by expand_site().
    "bom_core": [
        (
            "HID Signo-series readers — Corporate 1000 / HID 36-bit credential (STOP — confirm credential order with CEVA "
            "prior to manufacture)",
            19,
            485.00,
        ),
        (
            "Mercury LP1502 two-door controllers (preferred MP-/LP‑1502-class hardware per CEVA; MP‑1501-style single-door PoE-only "
            "controllers excluded)",
            7,
            1750.00,
        ),
        ("Mercury MR52 Supervised I/O (typical mix per panel)", 10, 395.00),
        (
            "LifeSafety Power (LSP) enclosures & dual-voltage distribution assemblies — "
            "12V controller rails + 24V locking hardware feeders per CEVA",
            7,
            2050.00,
        ),
        ("REX motion (DSC/Bosch-class — finalize SKU with Engineering)", 14, 88.00),
        ("Door contacts (interior / exterior blend)", 12, 42.00),
        ("Sounders / strobes (alarm + IN OUT door policy allowances)", 6, 240.00),
        ("Electric strikes / electrified hinge / actuator allowance (non-locksmith hardware)", 10, 175.00),
        (
            "Green + yellow structured patch cabling allowance — Cyber SSPS 270μm-class ACS headend/device patch tiers",
            1,
            1600.00,
        ),
        ("UPS batteries, PSU accessories & panel consumables (job lot)", 1, 1850.00),
        (
            "Conduit stubs + flex elbows targeting ~15 ft AFF at perimeter warehouse portals (coordinate ceiling "
            "penetrations with landlord/GC drawings)",
            1,
            395.00,
        ),
        (
            "J-hooks / strut / clamps for rafter transitions + masonry wall strapping allowances (raceway continuity per "
            "CEVA pathway detail)",
            1,
            525.00,
        ),
        ("Cable labels / numbering consumables / general wire-management hardware", 1, 180.00),
        ("Removal supplies, manifests, landfill & recycle allowance (existing wiring includes underground abandonment)", 1, 800.00),
        ("Freight consolidation / crane gate logistics allowance", 1, 850.00),
    ],
    # Wire model + lift (merged into full BOM by expand_site)
    "wire": {
        "avg_composite_homerun_ft": 300.0,  # planning avg per reader; must be ≤ MAX_RUN_FT
        "inout_homerun_count": 6,  # count of 22/6 IN/OUT homeruns (planning — confirm on drawings)
        "avg_inout_homerun_ft": 280.0,
    },
    "lift": {"rental_days": 6, "daily_all_in": 245.00},  # scissor/boom allowance incl. delivery — verify quote
}

SITE2 = {
    "id": "US-MJB-02",
    "name": "CL Warehouse (Building 3)",
    "address": "12008 Volunteer Blvd, Building 3, Mount Juliet, TN",
    "readers_bid": 27,
    "lp1502": 11,
    "phases": [
        ("A — Pre-mobilization", 25),
        ("B — Pro-Watch decommission", 100),
        ("C — Rough-in", 138),
        ("D — Headend install", 98),
        ("E — Device install", 96),
        ("F — Power-on checks", 24),
        ("G — ACM programming", 64),
        ("H — Commissioning / SAT & punch", 32),
        ("I — Closeout documentation", 50),
    ],
    "bom_core": [
        (
            "HID Signo-series readers — Corporate 1000 / HID 36-bit credential (STOP — confirm with CEVA prior to issuing "
            "PO)",
            27,
            485.00,
        ),
        (
            "Mercury LP1502 two-door controllers (CEVA-preferred Mercury MP-/LP‑1502 footprint; excludes MP‑1501-style "
            "single-door controllers)",
            11,
            1750.00,
        ),
        ("Mercury MR52 Supervised I/O boards (mix per clustered panel tier)", 14, 395.00),
        (
            "LifeSafety Power (LSP) enclosures — dual voltage 12 VDC controller + "
            "24 VDC electrified hardware distribution assemblies",
            11,
            2050.00,
        ),
        ("REX motion (finalize model sheet with security engineering)", 20, 88.00),
        ("Door position contacts (balanced interior / exterior ratios)", 16, 42.00),
        ("Sounders / strobes for alarmed-door & IN OUT policies", 10, 240.00),
        ("Electrified door hardware allowances (strike / actuator — non-locksmith)", 14, 175.00),
        (
            "Green + yellow ACS patch cabling + Cyber SSPS 270μm plant allowance (warehouse headend tiers)",
            1,
            2400.00,
        ),
        ("UPS / battery trays / misc power hardware (job lot)", 1, 2650.00),
        (
            "Conduit stubs to ~15 ft AFF at perimeter doors + stub seal hardware (coordinate with landlord structural "
            "drawings)",
            1,
            578.00,
        ),
        (
            "J-hook / seismic bracing allowances for attic & rafter cable routing transitions at warehouse entrances",
            1,
            798.00,
        ),
        ("Cable ID consumables / wire-management trim pieces", 1, 174.00),
        (
            "Decommission supplies incl. abandonment compliance for obsolete underground conductors (coordinate with EH&S)",
            1,
            950.00,
        ),
        ("Inbound freight consolidation / JIT staging allowance", 1, 1100.00),
    ],
    "wire": {
        "avg_composite_homerun_ft": 340.0,
        "inout_homerun_count": 8,
        "avg_inout_homerun_ft": 300.0,
    },
    "lift": {"rental_days": 9, "daily_all_in": 245.00},
}

# Louisville KY — CEVA table: 8 readers / 4 boards; bidding +5 reader buffer ⇒ 13 reader openings (verify on drawings).
SITE_LUI_02 = {
    "id": "US-LUI-02",
    "name": "CL Warehouse",
    "address": "5200 Interchange Drive, Louisville, KY 40229",
    "readers_bid": 13,
    "lp1502": 4,
    # KY state sales/use tax baseline — confirm local surcharge with CPA / customer (40229).
    "sales_tax_pct": 0.0600,
    "phases": [
        ("A — Pre-mobilization", 16),
        ("B — Pro-Watch decommission", 50),
        ("C — Rough-in (pathway, homeruns, headend rough-in)", 62),
        ("D — Headend install (LSP / LP1502 / MR52)", 45),
        ("E — Device install (readers, terminations, door interfaces)", 45),
        ("F — Power-on checks & comms troubleshooting", 14),
        ("G — ACM programming", 38),
        ("H — Commissioning / SAT & punch", 18),
        ("I — Closeout documentation", 32),
    ],
    "bom_core": [
        (
            "HID Signo-series readers — Corporate 1000 / HID 36-bit credential (ordering pause until CEVA credential "
            "workbook sign-off)",
            13,
            485.00,
        ),
        (
            "Mercury LP1502 two-door controllers (MP-/LP‑1502-tier hardware per CEVA; MP‑1501-class PoE single-door barred)",
            4,
            1750.00,
        ),
        ("Mercury MR52 I/O assemblies (balanced mix)", 6, 395.00),
        (
            "LifeSafety Power enclosures + dual-bus distribution (12V logic + 24V locking hardware feeders)",
            4,
            2050.00,
        ),
        ("REX motion heads (coordinate models with EH&S egress policy)", 10, 88.00),
        ("Door contacts", 8, 42.00),
        ("Visual / audible annunciators allowances", 4, 240.00),
        ("Electrified-door interface allowance (strike / actuator packages)", 6, 175.00),
        (
            "Cyber SSPS-compliant green/yellow ACS patch cabling + accessories (Cyber 270μm spec)",
            1,
            950.00,
        ),
        ("UPS, cabinet accessories & panel consumables (job lot)", 1, 1100.00),
        ("Conduit stub ups / flex elbows ~15 AFF at insulated warehouse entrances", 1, 259.00),
        (
            "J-hooks + wall / rafter fastening hardware for ACS vertical transitions on warehouse elevations",
            1,
            296.00,
        ),
        ("Cable ID / numbering / miscellaneous wire grooming", 1, 95.00),
        ("Rip-out supplies + debris handling (covers underground abandoned legs per CEVA)", 1, 550.00),
        ("Freight / trucking allowance — Louisville metro", 1, 950.00),
    ],
    "wire": {
        "avg_composite_homerun_ft": 280.0,
        "inout_homerun_count": 4,
        "avg_inout_homerun_ft": 270.0,
    },
    "lift": {"rental_days": 4, "daily_all_in": 245.00},
}

# Puerto Rico additional site — 9 ACS readers (+5 bidding buffer ⇒ 14); 3 CI/CO doors ⇒ 22/6 SHLD homerun count 6 pending survey.
SITE_PR_CAR_01 = {
    "id": "PR-CAR-01",
    "name": "Carolina — Julio N. Matos Industrial Park",
    "address": "Lot 26 Campeche St., Julio N. Matos Industrial Park, Martin Gonzalez Ward, Carolina, Puerto Rico 00983",
    "site_contact_note": (
        "CEVA site contact (addendum):\nJeannette M. Rexach · Station Manager – SJU\n"
        "jeannette.rexach@cevalogistics.com · Tel +1 (787) 253-3020 · Mobile +1 (787) 406-3204\n"
        "Additional site scope clarification: nine readers total; three read-in/read-out perimeter doors "
        "(CEVA cabling standards applied same as originating RFP)."
    ),
    "readers_bid": 14,
    "lp1502": 5,
    # PR SST/IVU & municipal nuances — refine with purchaser / CPA prior to contracting.
    "sales_tax_pct": 0.105,
    "phases": [
        ("A — Pre-mobilization", 18),
        ("B — Pro-Watch decommission / legacy cabling removal incl. underground (planning)", 56),
        ("C — Rough-in (raceway stubs ~15 AFF, straps/rafters, homeruns, headend)", 72),
        ("D — Headend install (LSP dual-voltage layout / LP1502 / MR52)", 54),
        ("E — Device install + field programming prep", 50),
        ("F — Power-on checks / comms", 17),
        ("G — ACM programming / IP worksheet execution", 44),
        ("H — Commissioning / SAT / Security Technology stakeholder walkthrough", 22),
        ("I — As-built turnover (CEVA-property deliverables)", 38),
    ],
    "bom_core": [
        (
            "HID Signo-series readers — Corporate 1000 / HID 36-bit credential (STOP — confirm credential order with "
            "CEVA Security Technology prior to release)",
            14,
            485.00,
        ),
        (
            "Mercury LP1502 two-door controllers (preferred MP-/LP‑1502 class hardware; Mercury MP‑1501-style "
            "single-door PoE-only controllers excluded per CEVA standard)",
            5,
            1750.00,
        ),
        ("Mercury MR52 Supervised I/O (typical mix per panel)", 7, 395.00),
        (
            "LifeSafety Power (LSP) enclosure & dual-voltage distribution — 12V controller bus + "
            "24V electrified hardware branches per CEVA",
            5,
            2050.00,
        ),
        ("REX motion (DSC / Bosch-class — finalize model with RFI)", 12, 88.00),
        ("Door contacts (interior / exterior mix)", 10, 42.00),
        ("Sounder/strobe allowances (alarm / IN OUT policy-driven)", 5, 240.00),
        ("Electric strike / EL / door-interface allowance (non-locksmith hardware)", 8, 175.00),
        (
            "Green + yellow Ethernet patch cabling — Cyber SSPS 270μm-class ACS patch plant allowance",
            1,
            1025.00,
        ),
        ("Battery plant, PSU accessories, miscellaneous panel hardware (job lot)", 1, 1200.00),
        (
            "Conduit stubs + flex to ~15 ft AFF at perimeter warehouse doors — planning lump (excludes civil/GC conduit "
            "not shown on ACS drawings)",
            1,
            392.00,
        ),
        (
            "J-hooks / rafter & wall-structure strapping allowances (raceway tier per CEVA warehouse doors)",
            1,
            428.00,
        ),
        ("Cable labeling consumables / wire-management hardware (remainder)", 1, 146.00),
        ("Decommission supplies, packaging, dumpster / recycle allowance", 1, 580.00),
        ("Freight / logistics / bonded cargo allowance — island-bound material", 1, 2625.00),
    ],
    "wire": {"avg_composite_homerun_ft": 287.0, "inout_homerun_count": 6, "avg_inout_homerun_ft": 272.0},
    "lift": {"rental_days": 6, "daily_all_in": 265.00},
}


def money(n: float) -> str:
    return f"${n:,.2f}"


def _feet_with_loop_and_waste(num_runs: int, avg_path_ft: float) -> float:
    """Billable cable feet for parallel homeruns (planning)."""
    per_run = avg_path_ft + SERVICE_LOOP_FT
    return float(num_runs) * per_run * (1.0 + WASTE_PCT)


def _bom_site_scale_multiplier(readers_bid: int) -> float:
    """Rough scale for provisional standards-driven allowances versus ~19-opening baseline."""
    if readers_bid >= 25:
        return 1.12
    if readers_bid >= 17:
        return 1.0
    return 0.93


def acm_license_pack_qty_ceva_double(readers_bid: int) -> int:
    """
    ACM 16-reader SKU count planning hook: ceil(N ÷16) packs per ACM server × 2 servers × 2
    for CEVA 'sell with double licenses' posture (finalize with distributor + STPM).
    """
    per_server_need = math.ceil(int(readers_bid) / 16.0)
    return max(1, per_server_need) * 2 * 2


def _ceva_standards_bom_addon(site: dict) -> list[tuple]:
    """Indicative materials QIS retains as vendor so we never silent-drop CEVA network / optics / MDF duties."""
    m = _bom_site_scale_multiplier(int(site["readers_bid"]))
    rid = site["id"]
    qty_lic = acm_license_pack_qty_ceva_double(site["readers_bid"])

    def sc(base: float) -> float:
        return round(base * m, -1)

    rows: list[tuple] = [
        (
            "MDF / IDF rack adjunct + structured cable-management kit (vendor supplies when retrofit standards mandate new "
            "racks per CEVA; coordinate cabinet elevation/power with facilities)",
            1,
            sc(2850.0),
            "Adjust after rack elevation drawings + Site Leadership / STPM direction.",
        ),
        (
            "Yellow Cat6 plenum horizontal + patch allowances — ACS field / IDF interconnect toward corporate switch "
            "demarc per CEVA",
            1,
            sc(795.0),
            "Terminated & tested lengths after Network Support port map.",
        ),
        (
            "Cisco core switch license placeholder (reference SKU LIC-C9300-24A-5Y 5YR — qty TBD, non-returnable)",
            1,
            sc(4550.0),
            "Network Support verifies exact switch qty & portal assignment before PO.",
        ),
        (
            "Cisco edge switch license placeholder (reference LIC-9300L-24E-5Y 5YR — qty TBD)",
            1,
            sc(3325.0),
            "",
        ),
        (
            "Industrial / hardened Cisco IE entitlement allowance (reference SW-CON-SNT-IE002SLM class — qty TBD)",
            1,
            sc(2185.0),
            "",
        ),
        (
            "Fiber / optic allowance — multimode optics + OM3 LC jumpers "
            "(refs GLC-SX-MMD ~1GE, SFP-10G-SR ~10GE — qty TBD)",
            1,
            sc(2595.0),
            "",
        ),
        (
            "12-strand OM3 armored retrofit backbone allowance + LC LIU hardware segment (references CEVA interior "
            "retrofit OM3 tier — extents TBD with Building 3 cabling survey)",
            1,
            sc(4650.0),
            "",
        ),
        (
            "Network Support coordination allowance (labels, workstation consumables for serial worksheet / portal "
            "documentation — optics handling fee)",
            1,
            round(575.0 * m, -1),
            "Labor for serial capture split into programming / commissioning scopes.",
        ),
        (
            f"Avigilon ACM 16-reader licensing units ×{qty_lic} (planning qty applies CEVA double-license doctrine on "
            "dual-server architecture — confirm SKU with distributor)",
            float(qty_lic),
            2985.0,
            "Server 1 + Server 2 (failover) — align reader maps with failover strategy.",
        ),
    ]
    if rid == "PR-CAR-01":
        rows.append(
            (
                "Puerto Rico import duties / surcharge contingency on optics + controller hardware pallets",
                1,
                2100.0,
                "",
            ),
        )
    return rows


def _bom_desc_uncertain_heuristic(desc: str) -> bool:
    """Rough-cut: BOM lines likely to swing with distributor quotes / takeoffs / rentals."""
    dl = desc.lower()
    needles = (
        "job lot",
        "allowance",
        "dumpster",
        "freight",
        "logistics allowance",
        "misc.",
        "decommission supplies",
        "rental",
        "model per rfi",
        "finalize",
        "typical mix",
        "until survey",
        "survey to confirm",
        "qty tbd",
        "placeholder",
        "reference sku",
        "coordinate",
        "contingency",
        "planning qty",
        "planning lump",
        "finalize sku",
        "dual-server",
        "optic",
        "backbone",
        "license",
        "import duties",
        "interconnect",
    )
    if any(k in dl for k in needles):
        return True
    return dl.startswith("yellow plenum")


def expand_site(site: dict) -> dict:
    """Return a copy of site with `bom` = core + computed cable + lift rental rows."""
    out = {**site}
    bom = list(site["bom_core"])
    bom.extend(_ceva_standards_bom_addon(site))
    w = site["wire"]
    readers = int(site["readers_bid"])
    avg_c = float(w["avg_composite_homerun_ft"])
    avg_i = float(w["avg_inout_homerun_ft"])
    n_inout = int(w["inout_homerun_count"])

    if avg_c > MAX_RUN_FT or avg_i > MAX_RUN_FT:
        raise ValueError(
            f"{site['id']}: avg homerun exceeds CEVA {MAX_RUN_FT:g} ft max — add IDF/aux or reduce avg."
        )

    ft_comp = _feet_with_loop_and_waste(readers, avg_c)
    ft_inout = _feet_with_loop_and_waste(n_inout, avg_i)

    note_500 = f"planning; each run ≤{MAX_RUN_FT:g} ft to panel/IDF per CEVA; survey to confirm"

    bom.append(
        (
            f"Yellow plenum composite — reader homeruns (est. {readers} runs × "
            f"({avg_c:g}+{SERVICE_LOOP_FT:g}) ft × (1+{WASTE_PCT:.0%}) waste)",
            round(ft_comp, 0),
            COMPOSITE_PLENUM_PER_FT,
            note_500,
        )
    )
    if n_inout > 0:
        bom.append(
            (
                f"Yellow plenum 22/6 SHLD — IN/OUT homeruns (est. {n_inout} runs × "
                f"({avg_i:g}+{SERVICE_LOOP_FT:g}) ft × (1+{WASTE_PCT:.0%}) waste)",
                round(ft_inout, 0),
                SHIELDED_22_6_PER_FT,
                note_500,
            )
        )

    cable_misc = 650.00
    if site["id"] == "US-MJB-04":
        cable_misc = 450.00
    elif site["id"] == "US-LUI-02":
        cable_misc = 380.00
    elif site["id"] == "PR-CAR-01":
        cable_misc = 720.00
    bom.append(
        (
            "Cable misc. (ties, transition, spare box, label stock for pulls)",
            1,
            cable_misc,
            "",
        )
    )

    lift = site["lift"]
    ld = int(lift["rental_days"])
    dr = float(lift["daily_all_in"])
    bom.append(
        (
            f"Aerial / scissor lift rental (est. {ld} days, all-in incl. delivery/pickup allowance)",
            ld,
            dr,
            "Verify Sunbelt/United/etc. quote; operator training/OSHA site rules apply",
        )
    )

    out["bom"] = bom
    return out


def bom_base_total(site: dict) -> float:
    return sum(qty * unit for (_, qty, unit, _, _) in _bom_rows(site))


def _bom_rows(site: dict):
    """Normalize BOM rows to (desc, qty, unit, notes, uncertain)."""
    rows: list[tuple[str, float, float, str, bool]] = []
    for row in site["bom"]:
        if len(row) == 5:
            desc, qty, unit, notes, uncertain = row
            rows.append((desc, float(qty), float(unit), str(notes), bool(uncertain)))
            continue
        if len(row) == 4:
            desc, qty, unit, fourth = row
            if isinstance(fourth, bool):
                rows.append((desc, float(qty), float(unit), "", fourth))
                continue
            notes = str(fourth)
            unc = _bom_desc_uncertain_heuristic(desc) or _bom_desc_uncertain_heuristic(notes)
            rows.append((desc, float(qty), float(unit), notes, unc))
            continue
        if len(row) == 3:
            desc, qty, unit = row
            rows.append((desc, float(qty), float(unit), "", _bom_desc_uncertain_heuristic(desc)))
            continue
        raise ValueError(f"Bad BOM row (len={len(row)}): {row!r}")
    return rows


def labor_direct_hours(site: dict) -> float:
    return sum(h for _, h in site["phases"])


def labor_financials(site: dict) -> dict:
    """Direct labor $, PM 12% of direct labor $, subtotal, then 25% margin on that subtotal."""
    h = labor_direct_hours(site)
    direct = h * LABOR_RATE
    pm = direct * PM_PCT
    pre_margin = direct + pm
    margin = pre_margin * MARGIN_PCT
    labor_sell = pre_margin + margin
    return {
        "hours": h,
        "labor_direct": direct,
        "pm": pm,
        "pre_margin": pre_margin,
        "labor_margin": margin,
        "labor_sell": labor_sell,
    }


def material_financials(site: dict) -> dict:
    base = bom_base_total(site)
    margin = base * MARGIN_PCT
    sell = base + margin
    return {"material_base": base, "material_margin": margin, "material_sell": sell}


def tax_bond_ocip(
    material_sell: float, pre_tax_subtotal: float, tax_pct: float | None = None
) -> dict:
    """Tax on materials only (placeholder). Bond & OCIP on pre-tax construction subtotal."""
    rate = TAX_MATERIAL_PCT if tax_pct is None else tax_pct
    tax = material_sell * rate
    bond = pre_tax_subtotal * BOND_PCT
    ocip = pre_tax_subtotal * OCIP_PCT
    return {"tax": tax, "bond": bond, "ocip": ocip}


def find_logo_path() -> Path | None:
    candidates = [
        ROOT / "public" / "branding" / "qis-logo.png",
        ROOT / "public" / "branding" / "qis-logo.jpg",
    ]
    for p in candidates:
        if p.is_file():
            return p
    return None


def set_cell_shading(cell, fill_hex: str) -> None:
    """fill_hex like '1B3A5F' (no #)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def style_table_header_row(table, row_index: int = 0) -> None:
    row = table.rows[row_index]
    for cell in row.cells:
        set_cell_shading(cell, FILL_HEADER)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.color.rgb = COLOR_WHITE
                run.font.bold = True


def style_subtotal_row(table, row_index: int) -> None:
    for cell in table.rows[row_index].cells:
        set_cell_shading(cell, FILL_SUBTLE)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def site_financial_totals(site: dict) -> dict:
    lf = labor_financials(site)
    mf = material_financials(site)
    pre_tax = lf["labor_sell"] + mf["material_sell"]
    tbo = tax_bond_ocip(
        mf["material_sell"], pre_tax, tax_pct=site.get("sales_tax_pct")
    )
    grand = pre_tax + tbo["tax"] + tbo["bond"] + tbo["ocip"]
    return {
        "grand": grand,
        "labor_sell": lf["labor_sell"],
        "material_sell": mf["material_sell"],
        "material_base": mf["material_base"],
        "pre_tax": pre_tax,
        "tax": tbo["tax"],
        "bond": tbo["bond"],
        "ocip": tbo["ocip"],
    }


def cover_page(doc: Document, fin1: dict, fin2: dict) -> None:
    doc.core_properties.title = "CEVA Access Control — Quality Installation Systems"
    doc.core_properties.subject = "Pro-Watch decommission + Avigilon ACM (two sites)"

    logo = find_logo_path()
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo), width=Inches(2.35))
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Your logo goes here: save as public/branding/qis-logo.png and re-run — we’ll drop it in automatically.\n")
        r.italic = True
        r.font.color.rgb = COLOR_SLATE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ACCESS CONTROL MODERNIZATION")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = COLOR_ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Proposal for CEVA Logistics")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = COLOR_NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "Pro-Watch decommission + Avigilon ACM / Mercury — two Mount Juliet facilities, priced honestly.\n"
    )
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_SLATE

    # Banner strip
    banner = doc.add_table(rows=1, cols=1)
    c0 = banner.rows[0].cells[0]
    set_cell_shading(c0, FILL_HEADER)
    c0.text = (
        f"  {SITE1['id']} + {SITE2['id']}  •  Planning investment (both sites): {money(fin1['grand'] + fin2['grand'])}  "
    )
    for paragraph in c0.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.color.rgb = COLOR_WHITE
            run.font.size = Pt(11)
            run.bold = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(COMPANY["legal_name"] + "\n")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = COLOR_NAVY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = [
        COMPANY["street"],
        COMPANY["city_st_zip"],
        f'{COMPANY["phone"]}  •  {COMPANY["email"]}',
    ]
    if COMPANY.get("website"):
        lines.append(COMPANY["website"])
    p.add_run("\n".join(lines) + "\n")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Prepared for: CEVA Logistics\n").bold = True
    p.add_run(f"Proposal date: {date.today().strftime('%B %d, %Y')}\n")
    exp = date.today() + timedelta(days=PROPOSAL_VALID_DAYS)
    p.add_run(
        f"Valid {PROPOSAL_VALID_DAYS} calendar days through {exp.strftime('%B %d, %Y')}.\n"
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("OCIP (quick read): ").bold = True
    p.add_run(
        "Owner-Controlled Insurance Program — sometimes the owner sponsors wrap coverage; sometimes there’s a small "
        "contractor fee. We show a conservative allowance you can zero out once CEVA confirms the program.\n"
    )


def executive_snapshot_page(doc: Document, fin1: dict, fin2: dict) -> None:
    add_heading(doc, "Executive snapshot", level=1)
    p = doc.add_paragraph()
    r = p.add_run("The one-minute version.")
    r.bold = True
    r.font.color.rgb = COLOR_ACCENT
    r.font.size = Pt(12)

    doc.add_paragraph(
        "This isn’t “swap badges and disappear.” It’s a controlled decommission of legacy Pro-Watch, a disciplined "
        "Mercury + Avigilon ACM installation, cable and pathway work that respects CEVA’s standards (yes—including "
        "the 500 ft homerun rule), and a closeout package you can actually hand to IT and Security without wincing."
    )

    add_heading(doc, "At a glance — planning totals", level=2)
    snap = doc.add_table(rows=1, cols=3)
    snap.style = "Table Grid"
    h = snap.rows[0].cells
    h[0].text = "Site"
    h[1].text = "Scope (headline)"
    h[2].text = "Planning total"
    style_table_header_row(snap, 0)

    def snap_row(site_id: str, name: str, addr: str, fin: dict):
        row = snap.add_row().cells
        row[0].text = f"{site_id}\n{name}"
        row[1].text = addr
        row[2].text = money(fin["grand"])

    snap_row(SITE1["id"], SITE1["name"], SITE1["address"], fin1)
    snap_row(SITE2["id"], SITE2["name"], SITE2["address"], fin2)
    rtot = snap.add_row().cells
    set_cell_shading(rtot[0], FILL_SUBTLE)
    set_cell_shading(rtot[1], FILL_SUBTLE)
    set_cell_shading(rtot[2], FILL_SUBTLE)
    rtot[0].text = "Program"
    rtot[1].text = "Both sites (sum of planning totals above)"
    rtot[2].text = money(fin1["grand"] + fin2["grand"])
    for c in rtot:
        for paragraph in c.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    add_heading(doc, "Why Quality Installation Systems", level=2)
    wins = [
        "We translate RFP language into **staffable hours** and a **BOM you can defend**—not a single mystery lump sum.",
        "Rough-in, headend landing, programming, SAT, and **documentation** are split so procurement sees where the money goes.",
        "Exclusions are deliberate: **gates / full-height entries / voice intercom ecosystems** stay in Option Packages "
        "unless CEVA merges them into base SOW—meanwhile **Cisco / optical / rack / ACM licensing** burdens stay on the "
        "vendor line items so nothing ‘quietly disappears’ from procurement math.",
        "We’ve already thought about **long buildings**: if drawings push past the 500 ft rule, we talk **IDF / aux** early—not after cable is on the truck.",
    ]
    for w in wins:
        doc.add_paragraph(w, style="List Bullet")

    p = doc.add_paragraph()
    p.add_run("How to read the rest of this document: ").bold = True
    p.add_run(
        "Each site has its own BOM (with cable + lift logic), labor buildup, and financial summary. "
        "Assumptions and exclusions follow—because boring fine print is what keeps projects out of court.\n"
    )
    r = p.add_run("P.S. ")
    r.italic = True
    r.bold = True
    p.add_run(
        "If you want this deck even sharper, send a vector PDF title block next time—phone photos of D-size sheets "
        "are heroic, but they don’t do your scale bar any favors."
    ).italic = True


def cover_page_single(doc: Document, site: dict, fin: dict) -> None:
    """Cover for one-site proposal (Louisville KY, etc.)."""
    doc.core_properties.title = "CEVA Access Control — Quality Installation Systems"
    doc.core_properties.subject = f"Pro-Watch decommission + Avigilon ACM — {site['id']}"

    logo = find_logo_path()
    if logo:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(logo), width=Inches(2.35))
        doc.add_paragraph()
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run("Your logo goes here: save as public/branding/qis-logo.png and re-run.\n")
        r.italic = True
        r.font.color.rgb = COLOR_SLATE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("ACCESS CONTROL MODERNIZATION")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = COLOR_ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Proposal for CEVA Logistics — single site")
    r.bold = True
    r.font.size = Pt(16)
    r.font.color.rgb = COLOR_NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        f"Site: {site['id']} — {site['name']}\n{site['address']}\n"
        "Pro-Watch decommission + Avigilon ACM / Mercury (planning-grade estimate).\n"
    )
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = COLOR_SLATE

    banner = doc.add_table(rows=1, cols=1)
    c0 = banner.rows[0].cells[0]
    set_cell_shading(c0, FILL_HEADER)
    c0.text = f"  {site['id']}  •  Planning investment: {money(fin['grand'])}  "
    for paragraph in c0.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.color.rgb = COLOR_WHITE
            run.font.size = Pt(11)
            run.bold = True

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(COMPANY["legal_name"] + "\n")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = COLOR_NAVY
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(
        f'{COMPANY["street"]}\n{COMPANY["city_st_zip"]}\n'
        f'{COMPANY["phone"]}  •  {COMPANY["email"]}\n'
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Prepared for: CEVA Logistics\n").bold = True
    p.add_run(f"Proposal date: {date.today().strftime('%B %d, %Y')}\n")
    exp = date.today() + timedelta(days=PROPOSAL_VALID_DAYS)
    p.add_run(
        f"Valid {PROPOSAL_VALID_DAYS} calendar days through {exp.strftime('%B %d, %Y')}.\n"
    )
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("OCIP (quick read): ").bold = True
    p.add_run(
        "Allowance shown on financial summary unless CEVA confirms owner-sponsored wrap with no contractor fee.\n"
    )


def executive_snapshot_single(doc: Document, site: dict, fin: dict) -> None:
    add_heading(doc, f"Executive snapshot — {site['id']}", level=1)
    doc.add_paragraph(
        "Single-site Pro-Watch decommission plus Avigilon ACM / Mercury / HID Signo deployment at "
        f"{site['address']}. BOM and allowances mirror CEVA published pathway, cabling, headend, network, optics, "
        "and commissioning responsibilities so procurement does not under-scope vendor-loaded materials."
    )
    add_heading(doc, "Planning totals", level=2)
    snap = doc.add_table(rows=3, cols=2)
    snap.style = "Table Grid"
    snap.rows[0].cells[0].text = "Field"
    snap.rows[0].cells[1].text = "Value"
    style_table_header_row(snap, 0)
    snap.rows[1].cells[0].text = "Site identity"
    snap.rows[1].cells[1].text = f"{site['id']} — {site['name']}"
    snap.rows[2].cells[0].text = "Planning total (incl. tax/bond/OCIP lines)"
    snap.rows[2].cells[1].text = money(fin["grand"])
    doc.add_paragraph()
    tp = doc.add_paragraph()
    plc = math.ceil(int(site["readers_bid"]) / 16.0)
    buy_qty = acm_license_pack_qty_ceva_double(site["readers_bid"])
    rl = tp.add_run(
        "ACM licensing posture: allocate "
        f"{buy_qty}× (16-door capacity SKU planning units) referencing CEVA’s dual-server + double-license doctrine "
        f"(baseline ceil({plc}) packs per ACM server ×2 servers ×2). Final counts released only after failover reader map "
        "and distributor quote reconciliation."
    )
    rl.font.color.rgb = COLOR_SLATE
    rl.italic = True



def ceva_vendor_commitments_page(doc: Document) -> None:
    """Cross-site mirror of written CEVA standards so omissions happen with eyes open, not by accident."""
    add_heading(doc, "CEVA access control standards → QIS commitments", level=1)
    ip = doc.add_paragraph()
    ir = ip.add_run(
        "Quality Installation Systems (QIS), as vendor, aligns this proposal material with CEVA’s published ACS "
        "standards—including vendor-owned procurement for structured cabling interconnects toward the enterprise "
        "switch, commissioning documentation obligations, optics-class hardware, MDF/IDF rack adjunct allowances, ACM "
        "licensing duplication rules, Cisco reference SKUs/SFP/MM jumpers—and will narrow exact counts after Network "
        "Support issues the IP / switch matrix. Scope items that genuinely await a separate subsystem design (heavy "
        "vehicle gates without readers, unrelated intercom package, etc.) stay in the Addon Options workbook unless the "
        "executed CEVA SOW explicitly folds them here."
    )
    ir.font.color.rgb = COLOR_SLATE
    ir.italic = True

    bl = [
        (
            "**1 — Door hardware / readers.** HID Signo platform; credential format Corporate 1000 / HID 36-bit with "
            "written CEVA approval before issuance; raceway stubs terminate approximately **15 ft AFF**."
        ),
        "**2 — Cabling.** Existing ACS cable (surface + abandoned underground conductors per survey) ripped out "
        "and replaced; yellow plenum composite homeruns; yellow plenum 22/6 shielded for CI/CO doors—each homerun "
        f"priced with **{SERVICE_LOOP_FT:g} ft** slack at reader + **≤{MAX_RUN_FT:g} ft** rule to MDF/headend absent "
        "approved IDF strategy.",
        (
            "**3 — Cable support.** J-hooks/straps along rafters and structure per warehouse detail—not implied only in "
            "labor—we keep discrete pathway allowances visible in BOM."
        ),
        (
            "**4 — Panels / enclosures.** Locate at legacy ProWatch footprint or MDF/IDF-approved positions; dual-voltage "
            "LifeSafety Power assemblies (12 V controllers + 24 V locking hardware feeders); Mercury LP1502-tier "
            "controllers only (explicitly exclude MP 1501 single-door controllers)."
        ),
        "**5 — Networking + optics.** Yellow Cat6 + patch interconnect fields run from ACS toward corporate switch "
        "demarc, plus Cisco SKU placeholders (five-year LIC references), multimode optics, OM3 LC jumpers, hardened "
        "IE entitlement—all vendor-procurement lines subject to quantity confirmation/non-returnability rules.",
        (
            "**6 — Retrofit OM3/backbone allowances.** Provision for armored OM3 strands + LC LIU segments when egress "
            "surveys dictate IDFs or supplemental fiber—not buried as 'someone else buys fiber.'"
        ),
        (
            "**7 — Programming & turnover.** Responsible for ACM controller door mapping, DHCP/IP worksheets, VLAN "
            "coordination gates, ACM browser-only UI alignment, failover licensing loads, serialization logs for optics, "
            "as-built deliverables (**CEVA property** upon acceptance), SAT with Security Technology or Regional delegate."
        ),
    ]
    for b in bl:
        doc.add_paragraph(b, style="List Bullet")


def assumptions_page(doc: Document, material_tax_pct: float | None = None):
    pct_display = TAX_MATERIAL_PCT if material_tax_pct is None else material_tax_pct
    add_heading(doc, "Assumptions & exclusions", level=1)
    p = doc.add_paragraph()
    r = p.add_run("The fine print that keeps everyone friends.")
    r.italic = True
    r.font.color.rgb = COLOR_ACCENT
    items = [
        "Labor hours are planning envelopes until engineered drawings finalize; mobilization durations still align with "
        "the June 25, 2026 program milestone communicated by CEVA.",
        (
            "**Vendor posture:** Quality Installation Systems (QIS), as contracted vendor, adopts CEVA’s Access Control "
            "Standards text for obligations called out therein—particularly removal/re-pull of ACS wiring (surface + buried), "
            "pathway build per warehouse raceway sketches, MDF/headend patching toward corporate switches, Cisco/SFP/MM "
            "reference SKUs, ACM dual-server licensing multiples, commissioning documentation turnover, optics serial "
            "capture, VPN coordination for remote support (duration per CEVA policy), rack adjuncts whenever standards "
            "mandate racks be supplied inside the modernization package."
        ),
        f"Loaded labor rate: ${LABOR_RATE:.2f}/hr; PM {PM_PCT:.0%} on direct labor dollars; contractor margin {MARGIN_PCT:.0%} "
        "on labor subtotal and on material BOM base thereafter.",
        f"Material BOM values are indicative planning dollars—replace with stamped distributor quotes. "
        f"Anything marked **qty TBD**/placeholder tracks **Network Support confirmations** plus CEVA’s non-returnable Cisco rules.",
        (
            "**Site electrical:** 120 VAC originates within roughly **10 ft** of enclosures per landlord practice; "
            "QIS extends feeders into LSP panels / buck-boost assemblies per drawing."
        ),
        (
            "Vehicle gate ironmongery/intercom-heavy packages absent reader counts remain in Addon Option Packages unless the "
            "executed purchase order merges them—the standards language still acknowledges QIS-managed ACS cable and "
            "integration where those peripherals interface."
        ),
        f"Sales & use / IVU placeholders use {pct_display:.2%} on taxable material sell (**verify SST/IVU on island sites**, "
        "metro surcharges inside Tennessee/Kentucky, etc.).",
        f"Surety bond {BOND_PCT:.2%} + OCIP {OCIP_PCT:.2%} allowances ride pre-tax totals until CEVA instructs removal.",
        f"Cable envelopes honor **≤{MAX_RUN_FT:g} ft** homeruns (+{SERVICE_LOOP_FT:g} ft slack +{WASTE_PCT:.0%} waste math). "
        "Verified exceedances initiate IDF/aux fiber BOM deltas via change notice.",
        "Lift rental lines fund aerial pathway reach for rafter transitions; OSHA trained operators billed inside labor scopes.",
        "Credential purchase orders pause until Corporate 1000 / HID 36-bit confirmation email from CEVA Security Technology logs.",
        "Addon Options workbook captures optional gate packages/turnstile/intercom scope when procurement wants clean separation.",
    ]
    for a in items:
        doc.add_paragraph(a, style="List Bullet")



def add_placeholder_bom_list(doc: Document, site: dict) -> None:
    """Separate roll-call of BOM lines flagged as quote / takeoff placeholders."""
    rows = [(d, q, u, n) for d, q, u, n, uc in _bom_rows(site) if uc]
    if not rows:
        return

    add_heading(
        doc,
        f"BOM planning / quote placeholders — {site['id']} (included in BOM total above)",
        level=3,
    )
    i = doc.add_paragraph()
    ir = i.add_run(
        "These lines remain in the main BOM rollup; summarized here because dollars swing most with distributor quotes, "
        "surveyed path lengths / footages, allowances, or lift invoices — not necessarily because the SKU model is unknown."
    )
    ir.italic = True
    ir.font.color.rgb = COLOR_SLATE

    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Item"
    hdr[1].text = "Qty"
    hdr[2].text = "Extended (plan)"
    hdr[3].text = "Field / quote driver"
    style_table_header_row(t, 0)

    subt = 0.0
    for desc, qty, unit, notes in rows:
        ext = float(qty) * float(unit)
        subt += ext
        rr = t.add_row().cells
        rr[0].text = desc
        qdisp = str(int(qty)) if float(qty) == int(qty) else str(qty)
        rr[1].text = qdisp
        rr[2].text = money(ext)
        rr[3].text = notes if notes else "—"

    sr = t.add_row().cells
    sr[0].text = "Subtotal (appendix rows only)"
    sr[1].text = "—"
    sr[2].text = money(subt)
    sr[3].text = "Full material totals include margin + tax on financial summary."
    style_subtotal_row(t, len(t.rows) - 1)


def add_bom_table(doc: Document, site: dict):
    add_heading(doc, f"Bill of materials (indicative market pricing) — {site['id']}", level=2)
    p = doc.add_paragraph()
    rp = p.add_run(
        "Quantities and unit costs are planning placeholders until distributor quotes and engineered drawings are issued; "
        "quote- and takeoff-sensitive lines are summarized again in **BOM planning / quote placeholders** immediately "
        "below — the appendix mirrors the rows above (**not additive**)."
    )
    rp.italic = True

    t = doc.add_table(rows=1, cols=5)
    t.style = "Table Grid"
    h = t.rows[0].cells
    h[0].text = "Item"
    h[1].text = "Qty"
    h[2].text = "Unit (est.)"
    h[3].text = "Extended (est.)"
    h[4].text = "Notes"
    style_table_header_row(t, 0)

    total = 0.0
    for desc, qty, unit, notes, _uncertain in _bom_rows(site):
        ext = float(qty) * float(unit)
        total += ext
        row = t.add_row().cells
        row[0].text = desc
        qdisp = str(int(qty)) if float(qty) == int(qty) else str(qty)
        row[1].text = qdisp
        row[2].text = money(unit)
        row[3].text = money(ext)
        row[4].text = notes

    row = t.add_row().cells
    row[0].text = "Material base (pre-margin)"
    row[1].text = "—"
    row[2].text = "—"
    row[3].text = money(total)
    row[4].text = "Sum of BOM"
    style_subtotal_row(t, len(t.rows) - 1)


def add_labor_table(doc: Document, site: dict):
    add_heading(doc, "Labor buildup", level=2)
    t = doc.add_table(rows=1, cols=4)
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    hdr[0].text = "Phase / scope bucket"
    hdr[1].text = "Hours"
    hdr[2].text = f"Rate ({money(LABOR_RATE)}/hr)"
    hdr[3].text = "Extended"
    style_table_header_row(t, 0)

    total_h = 0.0
    for label, hrs in site["phases"]:
        row = t.add_row().cells
        row[0].text = label
        row[1].text = str(hrs)
        row[2].text = money(LABOR_RATE)
        row[3].text = money(hrs * LABOR_RATE)
        total_h += hrs

    lf = labor_financials(site)
    row = t.add_row().cells
    row[0].text = "Direct labor subtotal"
    row[1].text = str(int(total_h)) if total_h == int(total_h) else str(total_h)
    row[2].text = "—"
    row[3].text = money(lf["labor_direct"])

    row = t.add_row().cells
    row[0].text = f"PM / supervision ({PM_PCT:.0%} of direct labor $)"
    row[1].text = "—"
    row[2].text = "—"
    row[3].text = money(lf["pm"])

    row = t.add_row().cells
    row[0].text = f"Subtotal before contractor margin (direct + PM)"
    row[1].text = "—"
    row[2].text = "—"
    row[3].text = money(lf["pre_margin"])

    row = t.add_row().cells
    row[0].text = f"Contractor margin on labor services ({MARGIN_PCT:.0%})"
    row[1].text = "—"
    row[2].text = "—"
    row[3].text = money(lf["labor_margin"])

    row = t.add_row().cells
    row[0].text = "Labor sell total"
    row[1].text = "—"
    row[2].text = "—"
    row[3].text = money(lf["labor_sell"])
    style_subtotal_row(t, len(t.rows) - 1)


def financial_summary_table(doc: Document, site: dict, title: str):
    add_heading(doc, title, level=2)
    fin = site_financial_totals(site)

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Category"
    tbl.rows[0].cells[1].text = "Amount (USD)"
    style_table_header_row(tbl, 0)

    def row(k: str, v: str):
        c = tbl.add_row().cells
        c[0].text = k
        c[1].text = v

    row("Labor sell total (incl. PM & margin)", money(fin["labor_sell"]))
    row("Material base (BOM, pre-margin)", money(fin["material_base"]))
    row(f"Contractor margin on materials ({MARGIN_PCT:.0%})", money(fin["material_base"] * MARGIN_PCT))
    row("Material sell total", money(fin["material_sell"]))
    row("Pre-tax construction subtotal (labor sell + material sell)", money(fin["pre_tax"]))
    _tp = site.get("sales_tax_pct", TAX_MATERIAL_PCT)
    row(
        f"Est. sales/use tax on materials ({_tp:.2%} — verify jurisdiction)",
        money(fin["tax"]),
    )
    row(
        f"Performance & payment bond allowance ({BOND_PCT:.2%} of pre-tax subtotal — verify)",
        money(fin["bond"]),
    )
    row(
        f"OCIP / wrap allowance ({OCIP_PCT:.2%} of pre-tax subtotal — verify or zero)",
        money(fin["ocip"]),
    )
    row("TOTAL (this site)", money(fin["grand"]))
    style_subtotal_row(tbl, len(tbl.rows) - 1)

    return fin


def signature_block(doc: Document):
    doc.add_paragraph()
    add_heading(doc, "Let’s build it right", level=2)
    p = doc.add_paragraph()
    p.add_run(
        "We’re ready to align on drawings, IDF strategy, and cutover windows the moment CEVA says “go.” "
        "Questions welcome — we actually answer the phone.\n\n"
    )
    p.add_run("Respectfully submitted,\n\n")
    p.add_run(f"{SIGNER_NAME}\n").bold = True
    p.add_run(f"{SIGNER_TITLE}\n")
    p.add_run(f"{COMPANY['legal_name']}\n")


def site_pages(doc: Document, site: dict):
    doc.add_page_break()
    add_heading(doc, f"{site['id']} — site package", level=1)
    p = doc.add_paragraph()
    r = p.add_run("Scope + numbers for this facility.")
    r.italic = True
    r.font.color.rgb = COLOR_ACCENT
    p = doc.add_paragraph()
    p.add_run(f"{site['name']}\n").bold = True
    p.add_run(f"{site['address']}\n")
    p.add_run(
        f"Planning quantities: {site['readers_bid']} reader openings (CEVA +5 buffer unless addendum narrows literal count), "
        f"{site['lp1502']} LP1502 positions.\n"
    )
    if site.get("site_contact_note"):
        doc.add_paragraph(site["site_contact_note"])

    add_bom_table(doc, site)
    add_placeholder_bom_list(doc, site)
    add_labor_table(doc, site)
    return financial_summary_table(
        doc,
        site,
        f"Financial summary — {site['id']}",
    )


def combined_page(doc: Document, s1: dict, s2: dict):
    doc.add_page_break()
    add_heading(doc, "Program roll-up (both sites)", level=1)
    p = doc.add_paragraph()
    r = p.add_run("One number for the conference call.")
    r.italic = True
    r.font.color.rgb = COLOR_ACCENT

    pre_tax = s1["pre_tax"] + s2["pre_tax"]
    tax = s1["tax"] + s2["tax"]
    bond_base = pre_tax
    bond = bond_base * BOND_PCT
    ocip = bond_base * OCIP_PCT
    grand = pre_tax + tax + bond + ocip

    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = "Table Grid"
    tbl.rows[0].cells[0].text = "Line"
    tbl.rows[0].cells[1].text = "Amount (USD)"
    style_table_header_row(tbl, 0)

    def row(k: str, v: str):
        c = tbl.add_row().cells
        c[0].text = k
        c[1].text = v

    row("US-MJB-04 — total (per site table)", money(s1["grand"]))
    row("US-MJB-02 — total (per site table)", money(s2["grand"]))
    row("—", "—")
    row("Recomputed pre-tax subtotal (labor + material, both sites)", money(pre_tax))
    row("Combined est. tax on materials (same method)", money(tax))
    row("Combined bond allowance (on combined pre-tax)", money(bond))
    row("Combined OCIP allowance (on combined pre-tax)", money(ocip))
    row("GRAND TOTAL (both sites, recomputed)", money(grand))
    style_subtotal_row(tbl, len(tbl.rows) - 1)

    p = doc.add_paragraph()
    p.add_run(
        "Pick one story and stick to it: either **sum of per-site totals** or the **recombined grand** above "
        "(they should be close; recombined applies bond/OCIP once on the combined pre-tax base)."
    ).italic = True


def build_louisville() -> None:
    doc = Document()
    e = expand_site(SITE_LUI_02)
    fin = site_financial_totals(e)
    tax_pct = SITE_LUI_02.get("sales_tax_pct")

    cover_page_single(doc, SITE_LUI_02, fin)
    doc.add_page_break()
    executive_snapshot_single(doc, SITE_LUI_02, fin)
    doc.add_page_break()
    assumptions_page(doc, material_tax_pct=tax_pct)
    doc.add_page_break()
    ceva_vendor_commitments_page(doc)
    doc.add_page_break()

    site_pages(doc, e)

    signature_block(doc)

    OUT_PATH_LUI02.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH_LUI02))
    print(f"Wrote: {OUT_PATH_LUI02}")


def build_pr_carolina() -> None:
    doc = Document()
    e = expand_site(SITE_PR_CAR_01)
    fin = site_financial_totals(e)
    tax_pct = SITE_PR_CAR_01.get("sales_tax_pct")

    cover_page_single(doc, SITE_PR_CAR_01, fin)
    doc.add_page_break()
    executive_snapshot_single(doc, SITE_PR_CAR_01, fin)
    doc.add_page_break()
    assumptions_page(doc, material_tax_pct=tax_pct)
    doc.add_page_break()
    ceva_vendor_commitments_page(doc)
    doc.add_page_break()

    site_pages(doc, e)

    signature_block(doc)

    OUT_PATH_PR_CAR01.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH_PR_CAR01))
    print(f"Wrote: {OUT_PATH_PR_CAR01}")


def build() -> None:
    doc = Document()
    e1 = expand_site(SITE1)
    e2 = expand_site(SITE2)
    fin1 = site_financial_totals(e1)
    fin2 = site_financial_totals(e2)

    cover_page(doc, fin1, fin2)
    doc.add_page_break()
    executive_snapshot_page(doc, fin1, fin2)
    doc.add_page_break()
    assumptions_page(doc)
    doc.add_page_break()
    ceva_vendor_commitments_page(doc)
    doc.add_page_break()

    s1 = site_pages(doc, e1)
    s2 = site_pages(doc, e2)
    combined_page(doc, s1, s2)
    signature_block(doc)

    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_PATH))
    print(f"Wrote: {OUT_PATH}")


def main() -> None:
    p = argparse.ArgumentParser(description="Generate CEVA ACM bid DOCX.")
    p.add_argument(
        "--site",
        choices=("TN_PAIR", "US-LUI-02", "PR-CAR-01"),
        default="TN_PAIR",
        help="TN_PAIR = Mount Juliet two-site; US-LUI-02 = Louisville; PR-CAR-01 = Carolina PR addendum site.",
    )
    args = p.parse_args()
    if args.site == "US-LUI-02":
        build_louisville()
    elif args.site == "PR-CAR-01":
        build_pr_carolina()
    else:
        build()


if __name__ == "__main__":
    main()
